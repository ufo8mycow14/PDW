from __future__ import annotations

import argparse
from datetime import datetime, timezone
import re
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# Design preset: compact_reference_guide.
# Named override: a4_operator_manual (A4 portrait, 0.75 inch margins).
# Full reference tables retain the preset's 9360 DXA width and 120 DXA indent.

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
CYAN = "00A6B2"
INK = "1C2630"
MUTED = "5A6772"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F8FC"
PALE_GOLD = "FFF4D6"
PALE_RED = "FDECEC"
PALE_GREEN = "EAF6EF"
WHITE = "FFFFFF"
BORDER = "C7D0DA"
RED = "9B1C1C"
GOLD = "7A5A00"
GREEN = "246B45"

FONT = "Calibri"
MONO_FONT = "Consolas"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

SCRIPT_DIR = Path(__file__).resolve().parent
REPOSITORY_ROOT = SCRIPT_DIR.parent.parent
SCREEN_DIR = REPOSITORY_ROOT / "docs" / "manual" / "screenshots"

FIGURE_PATHS = {
    "01-main-window-v5.png": SCREEN_DIR / "v5-ui-x64-main-printwindow.png",
    "02-settings-center-v5.png": SCREEN_DIR / "v5-settings-about-me-final.png",
    "03-settings-development-preview.png": SCREEN_DIR / "settings-consolidation-x64.settings.png",
    "04-backup-restore-v5.png": SCREEN_DIR / "settings-consolidation-x64.png",
    "05-screen-options-preview.png": SCREEN_DIR / "settings-consolidation-x64.screen-options.png",
    "06-capcode-routing-v5.png": SCREEN_DIR / "capcode-x64-routes-enabled.png",
    "07-data-outputs-v5.png": SCREEN_DIR / "ui-data-outputs.png",
}

LINKS = {
    "project": "https://www.discriminator.nl/pdw/",
    "github": "https://github.com/Discriminator/PDW",
    "apprise_api": "https://github.com/caronc/apprise-api",
    "apprise_services": "https://appriseit.com/services/",
    "apprise_ntfy": "https://appriseit.com/services/ntfy/",
    "apprise_pushover": "https://appriseit.com/services/pushover/",
    "pushover_signup": "https://pushover.net/signup",
    "pushover_api": "https://pushover.net/api",
    "ntfy_docs": "https://docs.ntfy.sh/",
    "ntfy_phone": "https://docs.ntfy.sh/subscribe/phone/",
    "smtp2go_signup": "https://www.smtp2go.com/signup/",
    "smtp2go_settings": "https://support.smtp2go.com/hc/en-gb/articles/223087627-SMTP-Settings",
    "pipedream": "https://pipedream.com/docs",
    "pipedream_http": "https://pipedream.com/docs/workflows/building-workflows/triggers",
    "dreamhost": "https://www.dreamhost.com/hosting/",
    "dreamhost_sftp": "https://help.dreamhost.com/hc/en-us/articles/115000675027-FTP-overview-and-credentials",
    "zadig": "https://zadig.akeo.ie/",
    "rtlsdr": "https://www.rtl-sdr.com/rtl-sdr-quick-start-guide/",
}


CHAPTERS = [
    ("Read this first", "read-this-first"),
    ("Install and first start", "install-and-first-start"),
    ("Interface tour", "interface-tour"),
    ("Choose a signal source", "choose-a-signal-source"),
    ("Decode and improve a signal", "decode-and-improve-a-signal"),
    ("Understand decoded messages", "understand-decoded-messages"),
    ("Build filters and alerts", "build-filters-and-alerts"),
    ("Log, copy, display and review", "log-copy-display-and-review"),
    ("Email and push notifications", "email-and-push-notifications"),
    ("Secure file transfer", "secure-file-transfer"),
    ("Publish to a website or webhook", "publish-to-a-website-or-webhook"),
    ("Data outputs and delivery health", "data-outputs-and-delivery-health"),
    ("Daily operation and troubleshooting", "daily-operation-and-troubleshooting"),
    ("Reference", "reference"),
]


# Filled after the first render pass. Values are physical PDF page numbers.
TOC_PAGES: dict[str, int] = {
    "read-this-first": 4,
    "install-and-first-start": 6,
    "interface-tour": 8,
    "choose-a-signal-source": 10,
    "decode-and-improve-a-signal": 14,
    "understand-decoded-messages": 17,
    "build-filters-and-alerts": 19,
    "log-copy-display-and-review": 23,
    "email-and-push-notifications": 25,
    "secure-file-transfer": 28,
    "publish-to-a-website-or-webhook": 30,
    "data-outputs-and-delivery-health": 32,
    "daily-operation-and-troubleshooting": 34,
    "reference": 37,
}


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER, size: int = 5) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA}, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, *, name: str = FONT, size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_external_hyperlink(paragraph, text: str, url: str, *, color: str = BLUE) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), NAVY)
    r_pr.append(color_node)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, value, end):
        run._r.append(node)


def add_numbering_definition(document: Document, *, ordered: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


class ManualBuilder:
    def __init__(self, output: Path):
        self.output = output
        self.document = Document()
        self.bookmark_id = 1
        self.figure_number = 0
        self.bullet_num_id = add_numbering_definition(self.document, ordered=False)
        self.number_num_id = add_numbering_definition(self.document, ordered=True)
        self.configure_document()

    def configure_document(self) -> None:
        section = self.document.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.header_distance = Inches(0.49)
        section.footer_distance = Inches(0.49)
        section.different_first_page_header_footer = True

        styles = self.document.styles
        normal = styles["Normal"]
        normal.font.name = FONT
        normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = RGBColor.from_string(INK)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.18

        for style_name, size, color, before, after in (
            ("Heading 1", 18, BLUE, 18, 10),
            ("Heading 2", 13.5, BLUE, 14, 7),
            ("Heading 3", 11.5, DARK_BLUE, 10, 5),
        ):
            style = styles[style_name]
            style.font.name = FONT
            style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style.font.bold = True
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.keep_together = True

        caption = styles["Caption"]
        caption.font.name = FONT
        caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        caption.font.size = Pt(8.5)
        caption.font.italic = True
        caption.font.color.rgb = RGBColor.from_string(MUTED)
        caption.paragraph_format.space_before = Pt(3)
        caption.paragraph_format.space_after = Pt(8)
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if "TOC Entry" not in styles:
            toc_style = styles.add_style("TOC Entry", WD_STYLE_TYPE.PARAGRAPH)
        else:
            toc_style = styles["TOC Entry"]
        toc_style.font.name = FONT
        toc_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        toc_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        toc_style.font.size = Pt(10.5)
        toc_style.paragraph_format.space_after = Pt(4)
        tab = toc_style.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.3),
            alignment=2,
            leader=1,
        )
        _ = tab

        self.set_running_header_footer(section)

    def set_running_header_footer(self, section) -> None:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run("PDW 4.1 Beta User Manual")
        set_run_font(run, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_before = Pt(0)
        run = paragraph.add_run("PDW 4.1 Beta  |  Page ")
        set_run_font(run, size=8.5, color=MUTED)
        add_page_field(paragraph)

    def add_title_line(self, text: str, *, size: float, color: str, bold: bool = True,
                       align=WD_ALIGN_PARAGRAPH.CENTER, after: float = 6) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(after)
        run = paragraph.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold)

    def add_body(self, text: str, *, bold_lead: str | None = None, italic: bool = False,
                 align=None, keep: bool = False) -> None:
        paragraph = self.document.add_paragraph()
        if align is not None:
            paragraph.alignment = align
        paragraph.paragraph_format.keep_together = keep
        if bold_lead and text.startswith(bold_lead):
            lead = paragraph.add_run(bold_lead)
            set_run_font(lead, bold=True)
            rest = paragraph.add_run(text[len(bold_lead):])
            set_run_font(rest, italic=italic)
        else:
            run = paragraph.add_run(text)
            set_run_font(run, italic=italic)

    def add_link_paragraph(self, lead: str, link_text: str, url: str, tail: str = "") -> None:
        paragraph = self.document.add_paragraph()
        if lead:
            set_run_font(paragraph.add_run(lead))
        add_external_hyperlink(paragraph, link_text, url)
        if tail:
            set_run_font(paragraph.add_run(tail))

    def add_bullet(self, text: str) -> None:
        paragraph = self.document.add_paragraph()
        apply_num(paragraph, self.bullet_num_id)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.18
        set_run_font(paragraph.add_run(text))

    def add_step(self, text: str) -> None:
        paragraph = self.document.add_paragraph()
        apply_num(paragraph, self.number_num_id)
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing = 1.18
        set_run_font(paragraph.add_run(text))

    def add_callout(self, label: str, text: str, *, kind: str = "note") -> None:
        palette = {
            "note": (PALE_BLUE, BLUE),
            "warning": (PALE_GOLD, GOLD),
            "danger": (PALE_RED, RED),
            "success": (PALE_GREEN, GREEN),
        }
        fill, accent = palette[kind]
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.14)
        paragraph.paragraph_format.right_indent = Inches(0.10)
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.keep_together = True
        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        p_pr.append(shd)
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:color"), accent)
        left.set(qn("w:space"), "7")
        borders.append(left)
        p_pr.append(borders)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, color=accent, bold=True)
        set_run_font(paragraph.add_run(text))

    def add_heading(self, text: str, level: int = 2) -> None:
        self.document.add_paragraph(text, style=f"Heading {level}")

    def add_chapter(self, text: str, anchor: str) -> None:
        paragraph = self.document.add_paragraph(text, style="Heading 1")
        paragraph.paragraph_format.page_break_before = True
        add_bookmark(paragraph, anchor, self.bookmark_id)
        self.bookmark_id += 1

    def add_table(self, headers: Sequence[str], rows: Iterable[Sequence[str]], widths: Sequence[int],
                  *, font_size: float = 8.7, header_fill: str = LIGHT_BLUE) -> None:
        table = self.document.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            set_cell_shading(cell, header_fill)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.keep_together = True
            run = paragraph.add_run(header)
            set_run_font(run, size=font_size, color=NAVY, bold=True)
        set_repeat_table_header(table.rows[0])
        for row_index, values in enumerate(rows):
            row = table.add_row()
            for column_index, value in enumerate(values):
                cell = row.cells[column_index]
                if row_index % 2 == 1:
                    set_cell_shading(cell, "FAFBFC")
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                run = paragraph.add_run(str(value))
                set_run_font(run, size=font_size, color=INK)
        set_table_geometry(table, widths)
        set_table_borders(table)
        self.document.add_paragraph().paragraph_format.space_after = Pt(1)

    def add_code_block(self, lines: Sequence[str]) -> None:
        for line in lines:
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.22)
            paragraph.paragraph_format.right_indent = Inches(0.12)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            p_pr = paragraph._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F6F7F8")
            p_pr.append(shd)
            run = paragraph.add_run(line)
            set_run_font(run, name=MONO_FONT, size=8.5, color=INK)
        self.document.add_paragraph().paragraph_format.space_after = Pt(1)

    def add_figure(self, filename: str, caption: str, alt_text: str, *, width: float = 6.35) -> None:
        path = FIGURE_PATHS.get(filename, SCREEN_DIR / filename)
        if not path.exists():
            raise FileNotFoundError(path)
        self.figure_number += 1
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run()
        inline = run.add_picture(str(path), width=Inches(width))
        doc_pr = inline._inline.docPr
        doc_pr.set("descr", alt_text)
        doc_pr.set("title", caption)
        caption_paragraph = self.document.add_paragraph(
            f"Figure {self.figure_number}. {caption}",
            style="Caption",
        )
        caption_paragraph.paragraph_format.keep_with_next = False

    def add_toc(self) -> None:
        self.document.add_page_break()
        heading = self.document.add_paragraph("Contents", style="Heading 1")
        add_bookmark(heading, "contents", self.bookmark_id)
        self.bookmark_id += 1
        self.add_body(
            "The page numbers below refer to the PDF edition. In Word, use the linked chapter names to jump directly to a section."
        )
        for index, (title, anchor) in enumerate(CHAPTERS, start=1):
            paragraph = self.document.add_paragraph(style="TOC Entry")
            add_internal_hyperlink(paragraph, f"{index}. {title}", anchor)
            paragraph.add_run("\t")
            page = TOC_PAGES.get(anchor)
            run = paragraph.add_run(str(page) if page else "--")
            set_run_font(run, size=10.5, color=MUTED)

    def build(self) -> None:
        self.build_cover()
        self.build_front_matter()
        self.add_toc()
        self.chapter_read_first()
        self.chapter_install()
        self.chapter_interface()
        self.chapter_signal_source()
        self.chapter_decode()
        self.chapter_messages()
        self.chapter_filters()
        self.chapter_logs_display()
        self.chapter_notifications()
        self.chapter_transfer()
        self.chapter_publishing()
        self.chapter_operations()
        self.chapter_reference()
        self.set_core_properties()

    def build_cover(self) -> None:
        self.document.add_paragraph().paragraph_format.space_after = Pt(12)
        self.add_title_line("OPERATOR GUIDE", size=10, color=CYAN, after=14)
        self.add_title_line("PDW 4.1 Beta", size=31, color=NAVY, after=2)
        self.add_title_line("Paging Decoder for Windows", size=16, color=DARK_BLUE, bold=False, after=10)
        self.add_title_line(
            "Installation, decoding, filtering, notification, secure transfer and web publishing",
            size=11,
            color=MUTED,
            bold=False,
            after=15,
        )
        self.add_figure(
            "01-main-window.png",
            "PDW 4.1 Beta task-oriented main window.",
            "PDW 4.1 Beta main window showing the menu bar, text toolbar, monitor pane, filtered-message pane and signal indicator.",
            width=6.45,
        )
        self.add_title_line("Edition 1.0  |  August 2026  |  Covers PDW v4.1.0 Beta", size=9.5, color=MUTED, after=4)
        self.add_title_line("Portable 32-bit Windows application for Windows 10 and Windows 11", size=9, color=MUTED, bold=False, after=3)
        self.add_title_line("All messages, accounts, hostnames and credentials shown in this manual are fictional training data.", size=8.5, color=RED, bold=True, after=0)

    def build_front_matter(self) -> None:
        self.document.add_page_break()
        self.add_heading("How to use this guide", 1)
        self.add_callout(
            "Training data",
            "Every pager address, message, aircraft registration, network identifier, hostname, account name, token and password in this guide is synthetic. Do not copy example secrets into a live installation.",
            kind="warning",
        )
        self.add_body(
            "This manual replaces the 2010 PDW v3.1 manual for everyday operation of the PDW v4.1.0 beta. It retains the original protocol, message-format, filter, logging and display guidance while updating the workflow for the current six-menu interface and the new receiver, notification, transfer and publishing features."
        )
        self.add_heading("Feature status used in this manual", 2)
        self.add_table(
            ["Status", "Meaning"],
            [
                ("Available in 4.1 Beta", "Present in the supplied PDW v4.1.0 Beta executable and interface."),
                ("Preserved legacy", "Existing behaviour intentionally retained for compatibility."),
                ("Future / not included", "Discussed only to prevent users mistaking roadmap work for a current feature."),
            ],
            [2100, 7260],
        )
        self.add_heading("Ten-minute first run", 2)
        for step in (
            "Extract the whole PDW folder to a writable location such as C:\\PDW-Training. Do not run the executable from inside a ZIP file.",
            "Start PDW v4.1.0 Beta.exe and open Settings > Settings.",
            "Choose Follow Windows, Light or Dark, then open Decoder and input > Input setup.",
            "Select local soundcard, serial input, RTL-TCP, direct RTL-SDR USB, or a recording replay. Use only a signal you are permitted to receive.",
            "Choose Monitor > Decoding mode and select POCSAG/FLEX, ACARS, MOBITEX or ERMES.",
            "Check the signal bar and receive-quality indication. Correct input level, gain, bandwidth and tuning before changing decoder thresholds.",
            "Open Filters > Manage filters and create one synthetic training filter.",
            "Enable logging only after choosing a folder with suitable privacy and retention controls.",
            "Configure email, Apprise, transfer or publishing only after creating the required third-party account or service.",
            "Send a configuration-only test, then confirm an unfiltered synthetic message does not trigger filtered-only outputs.",
        ):
            self.add_step(step)

    def chapter_read_first(self) -> None:
        self.add_chapter("1. Read this first", "read-this-first")
        self.add_heading("What PDW does", 2)
        self.add_body(
            "PDW is a native Windows signal decoder. Combined with a suitable receiver or interface, it can monitor supported POCSAG, FLEX, ACARS, MOBITEX and ERMES signals; show decoded records; filter selected records; log, alert and publish operator-approved outputs; and provide statistics and signal diagnostics."
        )
        self.add_table(
            ["Mode", "Rates / role", "Current boundary"],
            [
                ("POCSAG", "512, 1200 and 2400 baud", "Legacy parser retained; audio and serial sources supported."),
                ("FLEX", "1600, 3200 and 6400 baud", "Legacy parser retained; enhanced four-level audio adds B/D phase decisions beside A/C."),
                ("ERMES", "6250 baud", "Legacy parser retained."),
                ("ACARS", "Aircraft data messages", "Optional parity checking; database files can add labels and aircraft information."),
                ("MOBITEX", "8000 baud in normal use", "Base/mobile sync, network frame sync and message-type controls retained."),
            ],
            [1700, 2500, 5160],
        )
        self.add_heading("What is new in 4.1 Beta", 2)
        self.add_table(
            ["Area", "Available now", "Compatibility / safety boundary"],
            [
                ("Windows interface", "Six task-oriented menus, text toolbar, Settings hub, System/Light/Dark themes, DPI-aware dialogs.", "Decoder algorithms and existing configuration files are retained."),
                ("Startup", "Optional WindowTitle and Start PDW with Windows.", "Automatic startup waits five seconds; one instance per application folder."),
                ("Audio", "WinMM first, event-driven WASAPI fallback and bounded device-loss retry.", "Legacy unsigned 8-bit path remains available."),
                ("Radio", "Reconnecting RTL-TCP and direct RTL-SDR USB through a bundled 32-bit receiver pack.", "Optional receiver DLLs load only after selection."),
                ("Diagnostics", "WAV/SigMF record and replay, waveform, quality history, metrics and calibration.", "Previous live source is restored after replay."),
                ("Secure transfer", "FTP, explicit FTPS, implicit FTPS and password-authenticated SFTP.", "FTPS validates certificates; SFTP requires a SHA-256 host key."),
                ("Push", "Filtered-only Apprise delivery with test, background queue and privacy-safe default body.", "Apprise API is operated separately; secrets use Credential Manager."),
                ("Publishing", "Static JSON/JSONL/RSS/Atom/HTML and generic HTTPS webhooks.", "Disabled by default; jurisdiction acknowledgement and published-copy privacy controls are required."),
            ],
            [1750, 3800, 3810],
            font_size=8.2,
        )
        self.add_heading("Features still being added or held behind evidence gates", 2)
        self.add_table(
            ["Not a current feature", "Current position"],
            [
                ("Second independent decoder / soft-decision FEC", "Not enabled. Representative licensed or redacted recordings and before/after evidence are required first."),
                ("SoapySDR", "Not included. The narrower RTL-TCP and librtlsdr paths are the supported modern radio routes."),
                ("MQTT and publishing batches", "Possible future extensions. Current static files and HTTPS JSON webhook are complete without them."),
                ("64-bit PDW", "Not included. The executable remains 32-bit to retain legacy serial and slicer compatibility."),
            ],
            [3300, 6060],
        )
        self.add_callout(
            "Privacy and law",
            "PDW cannot determine whether reception, storage, forwarding or publication is lawful in your location. Check permissions, radio rules, privacy obligations and publication rules before use. Treat decoded content as potentially private.",
            kind="danger",
        )
        self.add_figure(
            "16-about.png",
            "About dialog for the supplied PDW v4.1.0 beta.",
            "PDW About dialog showing version 4.1.0 Beta, supported protocols, copyright credits, GNU GPL licence and contributor credit.",
            width=4.5,
        )

    def chapter_install(self) -> None:
        self.add_chapter("2. Install and first start", "install-and-first-start")
        self.add_heading("Requirements", 2)
        self.add_table(
            ["Item", "Requirement"],
            [
                ("Operating system", "Windows 10 or Windows 11."),
                ("Application", "PDW v4.1.0 Beta.exe and the folders supplied with it."),
                ("Architecture", "32-bit application; it runs on normal 64-bit Windows but receiver DLLs imported into PDW must be compatible 32-bit librtlsdr builds."),
                ("Input", "A Windows audio device, supported serial/RS232 interface, RTL-TCP server, compatible RTL-SDR USB receiver, or a supported WAV/SigMF replay file."),
                ("Network", "Required only for external notification, email, transfer, webhook or network receiver features."),
            ],
            [2200, 7160],
        )
        self.add_heading("Portable installation", 2)
        for step in (
            "Create a normal writable folder, for example C:\\PDW-Training.",
            "Extract the complete release into that folder. Keep Receivers, docs, PDW.INI, filters.ini and the executable together.",
            "If Windows shows a download warning, verify the release source and checksum before running it. Do not bypass a warning for an unverified copy.",
            "Start PDW v4.1.0 Beta.exe. The default source remains the legacy local input and publishing features remain off.",
            "Close PDW normally after the first configuration change so the INI files are written cleanly.",
        ):
            self.add_step(step)
        self.add_callout(
            "Do not mix releases",
            "Keep each portable PDW installation in its own folder. Copying only the new executable over an incomplete old folder can omit receiver packages, documentation or dependency files.",
            kind="warning",
        )
        self.add_heading("Upgrade from an older PDW folder", 2)
        self.add_bullet("Close the old PDW instance.")
        self.add_bullet("Copy PDW.INI, filters.ini, Wavfiles and any required database files to a private backup location.")
        self.add_bullet("Keep Logfiles and custom recordings according to your retention policy; they may contain decoded private content.")
        self.add_bullet("Extract the beta to a new folder, then copy only the settings and data you intentionally want to migrate.")
        self.add_bullet("Confirm input, filters, logging and output paths before enabling automation.")
        self.add_heading("Name separate instances", 2)
        self.add_body(
            "A separate folder can monitor a separate receiver or network. Add a WindowTitle value under [PDW] so the main window and tray tooltip identify the instance."
        )
        self.add_code_block(["[PDW]", "WindowTitle=Riverland Training Desk"])
        self.add_body(
            "The resulting title is Riverland Training Desk - PDW v4.1.0 Beta. A single folder still permits only one running instance."
        )
        self.add_heading("Start with Windows", 2)
        self.add_figure(
            "02-settings-hub.png",
            "Central Settings hub with appearance, startup, input and automation groups.",
            "PDW Settings dialog showing theme controls, Start PDW with Windows, decoder and input buttons, and connections and automation buttons.",
            width=6.3,
        )
        self.add_body(
            "Open Settings > Settings and select Start PDW with Windows. Automatic launches wait five seconds after sign-in to allow audio devices and services to settle. Clearing the same box disables the startup entry. Manual launches remain immediate."
        )
        self.add_heading("Important files and folders", 2)
        self.add_table(
            ["Path", "Purpose / privacy"],
            [
                ("PDW.INI", "Main settings. External passwords and tokens for modern features are not written here."),
                ("filters.ini", "Filter definitions, labels, hit counters and related options."),
                ("Wavfiles", "Optional WAV alerts, including address- and text-specific files."),
                ("Logfiles", "Operator logs. Protect, retain and delete according to policy."),
                ("Receivers", "Bundled and imported optional receiver packages plus the Zadig driver tool."),
                ("Published", "Default static publishing output. May contain transformed decoded events."),
                ("PublishQueue / DeadLetter", "Pending and repeatedly failed webhook events. Treat as sensitive."),
                ("FileTransfer.log / Apprise.log", "Sanitized operational results. Apprise.log excludes endpoints, credentials and decoded content."),
            ],
            [2600, 6760],
            font_size=8.4,
        )

    def chapter_interface(self) -> None:
        self.add_chapter("3. Interface tour", "interface-tour")
        self.add_figure(
            "01-main-window.png",
            "Main window in the Light appearance.",
            "PDW main window with File, Edit, Monitor, Filters, Settings and Help menus; text toolbar; monitor pane; filtered pane; and signal indicator.",
            width=6.5,
        )
        self.add_heading("Window areas", 2)
        self.add_table(
            ["Area", "Use"],
            [
                ("Menu bar", "All commands, organised as File, Edit, Monitor, Filters, Settings and Help."),
                ("Text toolbar", "One-click Log, Copy, Filters, Settings, Statistics, Pause, Clear and Mode actions."),
                ("Monitor pane", "Upper pane containing displayed decoded messages."),
                ("Filtered pane", "Lower pane containing messages that meet active non-reject filters and are not Monitor Only."),
                ("Divider", "Drag to change the relative pane sizes; F11 switches pane sizing."),
                ("Signal indicator", "Activity and receive-quality indication. It is not a calibrated RF power meter."),
            ],
            [2200, 7160],
        )
        self.add_heading("The six menus", 2)
        self.add_table(
            ["Menu", "Main tasks"],
            [
                ("File", "Open or close a log, save copied data, print copied data, exit."),
                ("Edit", "Copy selection, monitor window or filtered window; clear monitor panes."),
                ("Monitor", "Input setup, volume, decoder mode, statistics, pause or resume."),
                ("Filters", "Manage and reload filters, filter options, counters and command/file toggles."),
                ("Settings", "Appearance, decoder/input, email, push, transfer, publishing and character set."),
                ("Help", "User guide, diagnostics and About."),
            ],
            [1800, 7560],
        )
        self.add_heading("Appearance and sizing", 2)
        self.add_body(
            "Follow Windows changes with the system appearance. Light and Dark keep the chosen appearance. The current beta uses Segoe UI for modern dialogs, current Windows common controls and per-monitor DPI awareness. Legacy message colours and monospaced message fonts remain configurable."
        )
        self.add_callout(
            "Accessibility",
            "Use Settings > Appearance > Font for a readable monospaced message font, View and columns to remove unneeded fields, and Scrollback to balance history against performance. Keyboard shortcuts remain available when a mouse is inconvenient.",
            kind="note",
        )
        self.add_heading("Toolbar actions", 2)
        self.add_table(
            ["Button", "Action"],
            [
                ("Log", "Open or close the selected log file."),
                ("Copy", "Copy selected decoded text."),
                ("Filters", "Open the filter manager."),
                ("Settings", "Open the central Settings hub."),
                ("Statistics", "View decoder statistics."),
                ("Pause", "Pause or resume display/decoding operation as provided by PDW."),
                ("Clear", "Choose which monitor pane to clear."),
                ("Mode", "Change the current decoding mode."),
            ],
            [1800, 7560],
        )

    def chapter_signal_source(self) -> None:
        self.add_chapter("4. Choose a signal source", "choose-a-signal-source")
        self.add_heading("Source decision guide", 2)
        self.add_table(
            ["Source", "Use it when", "Important notes"],
            [
                ("Local audio", "A receiver supplies discriminator, line, tape, speaker or earphone audio.", "WinMM is tried first; WASAPI is an automatic fallback. Audio input is the simplest modern route for many receivers."),
                ("Serial slicer / RS232", "You have compatible legacy hardware and drivers.", "Two- and four-level slicer and RS232 behaviour is retained. Do not install obsolete VXD drivers on current Windows."),
                ("RTL-TCP", "An RTL-TCP-compatible receiver runs locally or on a trusted network host.", "PDW receives IQ, configures the tuner, demodulates NFM and reconnects. Decoded text is not sent to the server."),
                ("Direct RTL-SDR USB", "A compatible RTL2832U/V3/V4/V4L receiver is attached to this PC.", "Uses the bundled 32-bit standard pack. The exact USB interface may require WinUSB through Zadig."),
                ("Replay", "You are diagnosing a licensed, redacted or synthetic recording.", "Replays through the normal protocol routines and restores the previous live source afterward."),
            ],
            [1700, 3200, 4460],
            font_size=8.2,
        )
        self.add_heading("Legacy local audio and serial", 2)
        self.add_figure(
            "03-legacy-input-setup.png",
            "Legacy Input Setup remains available inside the modern interface.",
            "PDW Interface Setup dialog with Serial Port and Soundcard groups, soundcard configuration, sample rate and audio device selection.",
            width=4.3,
        )
        for step in (
            "Connect the receiver output or compatible serial interface before opening PDW.",
            "Open Monitor > Input setup or Settings > Decoder and input > Input setup.",
            "For audio, select Soundcard, choose the device and start with the closest source preset, normally Discriminator or Earphone.",
            "For serial, select the exact COM port, Slicer or RS232, and the required level/mode. Current COM numbers above 9 are supported.",
            "Select the matching decoder mode, tune to permitted activity and watch the signal indicator.",
            "Adjust receiver/input level so the signal is visible without persistent clipping. If the legacy format cannot open, WASAPI fallback is automatic.",
        ):
            self.add_step(step)
        self.add_callout(
            "Receiver modification",
            "A discriminator tap can improve decoding but requires receiver-specific technical work. Use a qualified technician or the correct service documentation; an incorrect modification can damage equipment or defeat electrical safety.",
            kind="warning",
        )
        self.add_heading("Direct RTL-SDR USB", 2)
        self.add_figure(
            "04b-direct-rtl-sdr.png",
            "Direct RTL-SDR USB source with the bundled receiver pack.",
            "Radio and Signal Sources dialog set to Direct RTL-SDR USB receiver, with receiver package, device, frequency, sample rates, gain, correction, bandwidth, recording and calibration controls.",
            width=6.0,
        )
        self.add_body(
            "The release includes a 32-bit standard receiver package for common RTL2832U receivers and RTL-SDR Blog V3, V4 and V4L devices. Open folder shows the package. Add receiver imports a trusted 32-bit rtlsdr.dll or librtlsdr.dll and neighbouring dependencies into Receivers\\Custom after architecture and API checks."
        )
        self.add_heading("Install WinUSB only when required", 3)
        self.add_link_paragraph("PDW includes Zadig 2.9 under Receivers\\Driver Tools. The authoritative download and usage page is ", "Zadig - USB driver installation made easy", LINKS["zadig"], ".")
        for step in (
            "Disconnect unrelated non-essential USB devices.",
            "Run Zadig as administrator and choose Options > List All Devices only if the RTL-SDR interface is not visible.",
            "Select the receiver's Bulk-In Interface 0, often shown as RTL2838UHIDIR or a similar RTL name. Confirm the USB ID against the device documentation.",
            "Select WinUSB, then install or replace the driver.",
            "Return to PDW, choose Direct RTL-SDR USB receiver, select the connected device and use Test selected source.",
        ):
            self.add_step(step)
        self.add_callout(
            "Driver warning",
            "Selecting a keyboard, mouse, storage device, hub or unrelated receiver in Zadig can replace that device's driver and stop it working. PDW never launches Zadig or changes drivers automatically.",
            kind="danger",
        )
        self.add_link_paragraph("For current vendor-specific V3/V4 setup guidance, see the ", "RTL-SDR Blog Windows quick start guide", LINKS["rtlsdr"], ".")
        self.add_heading("RTL-TCP network receiver", 2)
        self.add_figure(
            "04-radio-and-replay.png",
            "RTL-TCP example using realistic fictional training values.",
            "Radio and Signal Sources dialog set to RTL-TCP compatible network receiver, with fictional host decoder-gateway.example.net, port 1234, frequency 148337500 Hz, sample rates, gain, PPM correction and NFM bandwidth.",
            width=6.0,
        )
        self.add_table(
            ["Field", "Training example", "Meaning"],
            [
                ("RTL-TCP host", "decoder-gateway.example.net", "Hostname or IP of the trusted RTL-TCP listener."),
                ("Port", "1234", "Common RTL-TCP port; use the listener's actual port."),
                ("Frequency (Hz)", "148337500", "Whole-hertz tuner frequency. Use only an authorised target."),
                ("IQ sample rate", "1024000", "Samples per second received from RTL-TCP."),
                ("Audio rate", "48000", "Internal demodulated audio rate."),
                ("Automatic gain", "On", "Start here; use manual gain only when needed."),
                ("Correction (ppm)", "1", "Frequency correction for receiver oscillator error."),
                ("NFM BW (Hz)", "12500", "Adjustable from 5000 to 25000; start near the signal's occupied bandwidth."),
            ],
            [1900, 2300, 5160],
            font_size=8.2,
        )
        self.add_heading("Record, replay and calibrate", 2)
        self.add_bullet("WAV recording stores normalized decoder audio as 16-bit mono.")
        self.add_bullet("SigMF recording stores normalized rf32_le samples plus sample-rate metadata.")
        self.add_bullet("Replay accepts PCM8, PCM16 or float32 WAV and real float32 SigMF as documented by the beta.")
        self.add_bullet("Recording observes normalized samples without changing the decoder input. The in-memory diagnostic safety limit is 25 million samples.")
        self.add_bullet("Calibrate replay evaluates all 1,000 legacy threshold, centering and resync combinations, then asks before applying its signal-based suggestion to Custom configuration.")
        self.add_callout(
            "Test-data rule",
            "Use only synthetic, redacted or explicitly licensed recordings in bug reports, training packs or public repositories. Live pager traffic may contain private information.",
            kind="danger",
        )

    def chapter_decode(self) -> None:
        self.add_chapter("5. Decode and improve a signal", "decode-and-improve-a-signal")
        self.add_heading("Select the decoder mode", 2)
        self.add_body(
            "Use Monitor > Decoding mode or the Mode toolbar button. PDW monitors one family at a time: POCSAG/FLEX, ACARS, MOBITEX or ERMES. Selecting the wrong family can show activity without valid messages."
        )
        self.add_figure(
            "08-decoder-options.png",
            "Decoder options while monitoring POCSAG/FLEX.",
            "PDW Options dialog showing POCSAG rates and function handling, FLEX rates and instructions, Mobitex controls, ACARS parity and optional title-bar information.",
            width=5.9,
        )
        self.add_heading("POCSAG", 2)
        self.add_bullet("Enable POCSAG and the rates expected on the monitored system: 512, 1200 and/or 2400.")
        self.add_bullet("Decode function numbers as default only when the network uses the protocol function bits consistently.")
        self.add_bullet("Show both numeric and alphanumeric can help when the type cannot be inferred reliably, but it increases duplicate interpretations on screen.")
        self.add_heading("FLEX", 2)
        self.add_bullet("Enable the required 1600, 3200 and/or 6400 rates.")
        self.add_bullet("Short Instructions can be shown or converted when the network uses dynamic group calls.")
        self.add_bullet("FlexTIME appears only when detected. Review system-time implications before allowing any time correction.")
        self.add_heading("MOBITEX", 2)
        self.add_bullet("Choose Base or Mobile bit sync for downlink or uplink monitoring.")
        self.add_bullet("Start with Check Frame Sync when discovering the network identifier, then enter the confirmed Frame Sync for cleaner decoding.")
        self.add_bullet("Use the minimum character count and TEXT/DATA/HPDATA controls to reduce unreadable transparent traffic.")
        self.add_heading("ACARS and ERMES", 2)
        self.add_bullet("ACARS parity checking can suppress corrupted records but may hide traffic when input quality is poor. Compare results during setup.")
        self.add_bullet("ERMES uses its retained 6250-baud parser and is best matched to a suitable four-level input.")
        self.add_heading("Improve reception in the right order", 2)
        for step in (
            "Confirm the correct decoder mode and a known-active, permitted signal.",
            "Confirm the source is actually selected and Test selected source succeeds for RTL inputs.",
            "Correct frequency and PPM, then set a reasonable NFM bandwidth.",
            "Adjust antenna position and receiver gain. Avoid persistent clipping.",
            "For local audio, try the supplied Discriminator, Earphone, Speaker Out and Tape/Rec Out presets before Custom.",
            "Record a short synthetic or authorised sample and compare diagnostics.",
            "Use Calibrate replay only after the source, level and tuning are correct, and review the suggestion before applying it.",
        ):
            self.add_step(step)
        self.add_figure(
            "09-general-options.png",
            "General behaviour settings for duplicate blocking, grouping, dates and log paths.",
            "PDW General Options dialog with duplicate blocking, linefeed conversion, message separation, date formats, confirm exit and logfile path controls.",
            width=5.7,
        )
        self.add_heading("Duplicate blocking", 2)
        self.add_body(
            "Duplicate blocking can compare the previous address and message, the previous message only, or address/function/message within a timer. Enable blocked.txt during tuning if you need to review what is being suppressed. A blocked duplicate is also excluded from filtered-only push routing."
        )
        self.add_heading("Signal indicator and new diagnostics", 2)
        self.add_body(
            "The signal indicator reacts to input activity and valid decoding; it is not calibrated RF signal strength. The Radio and Signal Sources dialog adds level, noise, clipping, eye opening, signal score, corrected versus uncorrectable error totals, FLEX A-D phase totals and a short receive-quality history. Use the combination, not one number, to judge changes."
        )
        self.add_heading("Enhanced decoding boundary", 2)
        self.add_body(
            "The 4.1 beta runs adaptive DC and envelope tracking beside the legacy slicer. Four-level audio FLEX preserves the legacy sign decisions for phases A/C while adding inner/outer decisions for phases B/D; low-confidence input falls back to the legacy 0/3 path. The hard-decision protocol parser and serial paths remain authoritative. A second full parser and soft FEC are not present."
        )

    def chapter_messages(self) -> None:
        self.add_chapter("6. Understand decoded messages", "understand-decoded-messages")
        self.add_callout(
            "Synthetic examples",
            "The records in this chapter are formatted like PDW output but are invented. Addresses, registrations, flights, times and content do not identify a real person, aircraft or service.",
            kind="note",
        )
        self.add_heading("Paging format: POCSAG, FLEX and ERMES", 2)
        self.add_table(
            ["Address", "Time", "Date", "Mode", "Type", "Rate", "Message"],
            [
                ("1234567", "14:13:33", "10-08-26", "POCSAG-1", "ALPHA", "1200", "TRAINING DRILL - Unit 12 report to base"),
                ("7654321", "14:15:08", "10-08-26", "POCSAG-2", "NUMERIC", "512", "0400000000"),
                ("202345678", "14:17:25", "10-08-26", "FLEX-A", "ALPHA", "1600", "TEST ONLY - Airport operations standby"),
                ("3456789", "14:20:41", "10-08-26", "ERMES-2", "ALPHA", "6250", "TRAINING WEATHER - exercise complete"),
            ],
            [1050, 900, 950, 1250, 900, 700, 3610],
            font_size=7.6,
        )
        self.add_table(
            ["Field", "Meaning"],
            [
                ("Address", "Pager or receiver address, normally seven or nine digits in the supported formats. Filters can match it exactly or with ? wildcards."),
                ("Time / Date", "The receive timestamp using the computer clock unless a supported, detected network time option is deliberately used."),
                ("Mode", "Protocol plus POCSAG function number, FLEX phase letter or ERMES designation."),
                ("Type", "Alphanumeric, Numeric, Tone-only, Binary, Secure or Transparent as supported by the protocol."),
                ("Rate", "The detected bit rate."),
                ("Message", "Decoded content. Faulty characters, types and matched text may use configured colours."),
            ],
            [1900, 7460],
        )
        self.add_heading("ACARS format", 2)
        self.add_table(
            ["Air.Reg", "Time", "Date", "Msg.No", "DBI", "Mode", "Message"],
            [
                ("VH-TST", "16:42:07", "10-08-26", "F18A", "4", "R", "H1 TRAINING DATA - DO NOT ACTION"),
            ],
            [1100, 900, 950, 850, 600, 650, 4310],
            font_size=8.0,
        )
        self.add_table(
            ["ACARS item", "Meaning"],
            [
                ("Aircraft registration", "Registration reported in the ACARS record. Database enrichment may add country and aircraft type."),
                ("Message number", "Four-character sequence indicator used by the network."),
                ("DBI", "Downlink Block Identifier used to identify retransmissions."),
                ("Mode", "Ground System Interface Configuration value."),
                ("Message label", "Two-character label; label.df can add or update descriptions."),
                ("Flight / airline / station", "Derived fields when the message and local database files provide enough information."),
                ("ACK / NAK", "Technical acknowledgement or not-acknowledged indication."),
            ],
            [2300, 7060],
        )
        self.add_heading("MOBITEX format", 2)
        self.add_table(
            ["MAN", "Time", "Date", "Sender", "Type", "Rate", "Message"],
            [
                ("7654321", "15:14:13", "10-08-26", "7001200", "TEXT", "8000", "TRAINING NETWORK TEST"),
                ("1234567", "15:16:41", "10-08-26", "7001201", "HPDATA", "8000", "B+&(AB+&6 TEST"),
            ],
            [1050, 900, 950, 1050, 900, 700, 3810],
            font_size=7.8,
        )
        self.add_table(
            ["MOBITEX item", "Meaning"],
            [
                ("MAN", "Seven-digit Mobitex Access Number for the receiving radio modem."),
                ("Sender", "Seven-digit sending MAN."),
                ("Type", "TEXT, DATA, HPDATA or STATUS."),
                ("Rate", "Normally 8000 in this PDW workflow."),
                ("Message", "Raw 8-bit payload. Efficient application data can look random without being encrypted; the end-user application defines its meaning."),
            ],
            [2300, 7060],
        )
        self.add_heading("Choose visible columns", 2)
        self.add_figure(
            "10-view-and-columns.png",
            "View and columns for the paging monitor.",
            "PDW Screen Options dialog showing seven configurable columns and FLEX group mode controls.",
            width=5.7,
        )
        self.add_body(
            "Settings > Appearance > View and columns controls which fields appear and their order. FLEX Group Mode can combine group messaging where appropriate. Keep enough identifiers to interpret and audit records, but hide unused fields for a clearer display."
        )
        self.add_heading("Character sets", 2)
        self.add_body(
            "Settings > Character set chooses Default ASCII or a character set loaded from language.df, such as German, Hebrew, Skyper or Swedish where supplied. Editing language.df adds entries after restart. Back up the file before editing and keep a known-good encoding."
        )

    def chapter_filters(self) -> None:
        self.add_chapter("7. Build filters and alerts", "build-filters-and-alerts")
        self.add_heading("What a filter does", 2)
        self.add_body(
            "A filter compares an address and/or message text. A normal match can place a record in the filtered pane, add a label, play a WAV, write separate files, increment a hit counter and, when separately enabled, route email or Apprise. Reject filters hide matching records. Monitor Only matches can label or sound without entering the lower pane."
        )
        self.add_figure(
            "05-filter-manager.png",
            "Filter manager with three fictional training filters.",
            "PDW Filters dialog with fictional POCSAG and FLEX filters for Regional Dispatch, Airport Operations and a Weather Drill.",
            width=6.5,
        )
        self.add_heading("Create a filter", 2)
        for step in (
            "Open Filters > Manage filters and choose Add.",
            "Choose FLEX Riccode, POCSAG Riccode, Text, ERMES, ACARS or MOBITEX as applicable.",
            "Enter the address or text rule. Use ? as a single-position address wildcard.",
            "For a POCSAG address, choose All or a specific function number.",
            "Add up to ten required text fragments separated by &, for example TRAINING&UNIT 12. Matching is not case-sensitive.",
            "Add a clear label and select Show Filter Label if operators need it in output.",
            "Choose Reject, Monitor Only, email, command file, sound and separate-file actions deliberately.",
            "Choose OK, then use a synthetic input to confirm both a matching and a non-matching record.",
        ):
            self.add_step(step)
        self.add_figure(
            "06-add-edit-filter.png",
            "POCSAG filter using a realistic fictional address and text rule.",
            "PDW Add Filter dialog with POCSAG Riccode 1234567, All functions, text TRAINING and UNIT 12, Regional Dispatch label, default sound and disabled email action.",
            width=3.8,
        )
        self.add_table(
            ["Field", "Behaviour"],
            [
                ("Reject", "Prevents matching messages from being displayed or routed as normal filtered output."),
                ("Filter type", "Selects address/message interpretation for the active protocol."),
                ("Address", "Exact address or ? wildcard pattern, such as 12?????."),
                ("Text", "Case-insensitive phrase, exact-message rule or up to ten required fragments separated by &."),
                ("Label", "Operator-facing description, optionally displayed and colour-coded."),
                ("Monitor Only", "Allows a match action without placing the message in the lower filtered pane; it is not sent through filtered-only Apprise."),
                ("Send email", "Active only when email is configured; use Selected filters only when each filter should decide."),
                ("Command file", "Runs the configured external command for this filter. Treat decoded text as untrusted input."),
                ("Separate files", "Writes up to three filter-specific files when filter-file output is enabled."),
                ("Hit counter", "Tracks count and last hit; reset individual or selected/all counters as required."),
            ],
            [2250, 7110],
            font_size=8.3,
        )
        self.add_heading("Worked training rules", 2)
        self.add_table(
            ["Goal", "Type / address", "Text", "Expected result"],
            [
                ("One POCSAG address plus two words", "POCSAG 1234567 / All", "TRAINING&UNIT 12", "Both fragments must appear in any case."),
                ("Address range", "POCSAG 12?????", "blank", "Any seven-digit address beginning 12."),
                ("Exact test phrase", "Text", "TRAINING COMPLETE", "With Match EXACT, longer messages do not match."),
                ("Suppress known test source", "Appropriate address", "TEST ONLY", "Reject hides matching records."),
            ],
            [2400, 2200, 1900, 2860],
            font_size=8.2,
        )
        self.add_heading("Filter options", 2)
        self.add_figure(
            "07-filter-options.png",
            "Filter file, display, sound and command options.",
            "PDW Filter Options dialog with filter file output, logged columns, descriptions, command file, message-type controls, sounds and default filter type.",
            width=5.9,
        )
        self.add_body(
            "Filter Options controls a shared filtered logfile, date-based names, logged columns, description placement, message types, filter-window behaviour, sounds and the default filter type. Show filtered messages only in filter window/log files removes them from the upper pane but does not change the source record itself."
        )
        self.add_heading("WAV alerts", 2)
        self.add_body(
            "The Default sound uses the chosen built-in WAV. Address-specific names can override it: 1234567.wav, or 1234567-fire.wav when address and text match. More specific address files take priority over wildcard names, which take priority over generic filtered and Monitor_only WAV files. Test volume outside a live incident workflow."
        )
        self.add_heading("Command-file placeholders", 2)
        self.add_table(
            ["Token", "Value"],
            [
                ("%1", "Address"), ("%2", "Time"), ("%3", "Date"), ("%4", "Mode"),
                ("%5", "Type"), ("%6", "Bitrate"), ("%7", "Message"), ("%8", "Label"),
            ],
            [1400, 7960],
        )
        self.add_callout(
            "Command safety",
            "Decoded message text can contain unexpected characters. Quote arguments, send them to a fixed trusted program, avoid shell interpretation where possible, and never build privileged commands from pager content.",
            kind="danger",
        )

    def chapter_logs_display(self) -> None:
        self.add_chapter("8. Log, copy, display and review", "log-copy-display-and-review")
        self.add_heading("Main logfile", 2)
        self.add_figure(
            "15-logfile.png",
            "Logfile dialog with date-based naming and selected columns.",
            "PDW Logfile dialog showing Enable logfile, example filename 260810.log, Use date as filename and seven log column checkboxes.",
            width=3.6,
        )
        for step in (
            "Open File > Log file or select Log on the toolbar.",
            "Choose a protected output folder under Settings > General behaviour.",
            "Enable the logfile, browse to a fixed file or select Use date as filename.",
            "Select the minimum columns required for operations and audit.",
            "Confirm the date format and write permissions, then inspect the first record.",
        ):
            self.add_step(step)
        self.add_callout(
            "Retention",
            "A logfile can contain addresses, timestamps and message content. Limit access, back it up only when required, and define a deletion schedule. Sanitized transfer/notification logs are not a substitute for protecting the decoded log itself.",
            kind="danger",
        )
        self.add_heading("Copy, save and print", 2)
        self.add_table(
            ["Command", "Result"],
            [
                ("Copy Selection", "Copies highlighted decoded text."),
                ("Copy monitor window", "Copies all text in the upper pane."),
                ("Copy filtered window", "Copies all text in the lower pane."),
                ("Save copied data", "Writes the current clipboard copy to a text file."),
                ("Print copied data", "Sends the current clipboard copy to the selected printer."),
            ],
            [2600, 6760],
        )
        self.add_heading("Display and operator comfort", 2)
        self.add_figure(
            "10-view-and-columns.png",
            "Columns and FLEX group display controls.",
            "PDW Screen Options with configurable visible columns and FLEX group mode controls.",
            width=5.7,
        )
        self.add_bullet("Colours opens a protocol-specific colour dialog for fields, filter matches, labels and bit errors.")
        self.add_bullet("Font accepts a non-proportional message font so columns remain aligned.")
        self.add_bullet("Scrollback controls monitor and filtered line history, mouse-wheel speed and pane sizing.")
        self.add_bullet("System tray can keep PDW running in the background and restore on configured message events.")
        self.add_bullet("Statistics shows recent message totals; use it with signal diagnostics rather than as a standalone quality measure.")
        self.add_bullet("Clear lets the operator clear either pane without deleting existing files already written to disk.")
        self.add_heading("Generated output inventory", 2)
        self.add_table(
            ["Output", "Created by", "Review point"],
            [
                ("Date or fixed .log", "Main logfile", "May contain all displayed messages."),
                ("Filter file(s)", "Filter options / per-filter files", "May contain selected private messages and labels."),
                ("blocked.txt", "Duplicate blocking review", "Use temporarily and inspect before deletion."),
                ("WAV / SigMF", "Diagnostic recording", "Signal data can still carry private content."),
                ("messages.* / index.html", "Static publishing", "Contains the published-copy fields selected by the operator."),
                ("FileTransfer.log", "Continuous transfer", "Operational results; password is excluded."),
                ("Apprise.log", "Push delivery", "Sanitized event/result entries; decoded text and destinations are excluded."),
            ],
            [2200, 2600, 4560],
            font_size=8.2,
        )

    def chapter_notifications(self) -> None:
        self.add_chapter("9. Email and push notifications", "email-and-push-notifications")
        self.add_callout(
            "Third-party accounts",
            "PDW does not create email, Apprise, ntfy or Pushover accounts. Set up and test the external service first, then enter only service-specific credentials in PDW. Terms, pricing and limits can change.",
            kind="note",
        )
        self.add_heading("Email with an SMTP relay", 2)
        self.add_figure(
            "11-email-settings.png",
            "Filtered email example using SMTP2GO settings and a fictional account.",
            "SMTP email dialog with mail.smtp2go.com, port 465, SSL and authentication enabled, and fictional example.net sender, recipient and SMTP username.",
            width=5.8,
        )
        self.add_heading("SMTP2GO setup example", 3)
        self.add_link_paragraph("Create an account at ", "SMTP2GO signup", LINKS["smtp2go_signup"], ", then create a dedicated SMTP user under Sending > SMTP Users.")
        for step in (
            "Verify the sender domain or sender address in the provider dashboard.",
            "In PDW open Settings > Connections and automation > Email.",
            "Select Enable Email Notification and choose All messages, Filtered messages, Filtered & Monitor-Only messages, or Selected filters only.",
            "Enter mail.smtp2go.com, port 465, select SSL, enable authentication, and enter the SMTP user name and SMTP password - not the website account password.",
            "Enter a verified From address and one or more intended To addresses.",
            "Select the minimum subject/body fields required and choose Test.",
            "For Selected filters only, edit each intended filter and select Send email.",
        ):
            self.add_step(step)
        self.add_link_paragraph("The provider's current host, port and encryption combinations are documented in ", "SMTP2GO SMTP Settings", LINKS["smtp2go_settings"], ".")
        self.add_callout(
            "Email privacy",
            "Email can leave the controlled network, remain in mailboxes and appear in notifications. Prefer filtered messages, minimise fields and use an organisation-approved relay and recipient list.",
            kind="warning",
        )
        self.add_heading("Apprise filtered-message push", 2)
        self.add_figure(
            "12-apprise-notifications.png",
            "Apprise dialog with fictional masked API settings.",
            "Apprise Notifications dialog showing Enable Apprise, masked API URL, username, password and destinations; the Include decoded text option is off.",
            width=6.2,
        )
        self.add_body(
            "Apprise is independent of email. Every non-rejected, non-duplicate message that reaches the lower filtered output can be queued for Apprise. Unfiltered and Monitor Only records are not pushed. By default the body says only A filtered message was received. Filter details and decoded text require explicit opt-in and may appear on a phone lock screen."
        )
        self.add_heading("Operate an Apprise API", 3)
        self.add_link_paragraph("PDW expects an operator-managed Apprise API. Deployment instructions and the container are maintained at ", "the official Apprise API project", LINKS["apprise_api"], ".")
        self.add_bullet("The beta client contract is Apprise API 1.5.1 with Apprise 1.12.0. Pin those versions or regression-test a newer release before changing it.")
        self.add_bullet("Place the API behind HTTPS and a reverse proxy that enforces Basic authentication. Do not expose the default container port to an untrusted network.")
        self.add_bullet("Use /notify for stateless delivery and supply one or more Destination URLs in PDW.")
        self.add_bullet("Use /notify/{key} for destinations stored on the server and leave PDW Destination URLs blank.")
        self.add_bullet("The endpoint, destinations and API credentials are stored in Windows Credential Manager and masked in PDW.")
        self.add_heading("Pushover phone setup", 3)
        self.add_link_paragraph("Create a user account at ", "Pushover signup", LINKS["pushover_signup"], ", install the phone app, and register a PDW application/API token in the dashboard.")
        self.add_body("Use the account User Key and application token in the Apprise destination format:")
        self.add_code_block(["pover://USER_KEY@APPLICATION_TOKEN"])
        self.add_link_paragraph("Pushover's account keys and application-token process are described in the ", "official Pushover Message API guide", LINKS["pushover_api"], ".")
        self.add_heading("ntfy phone setup", 3)
        self.add_link_paragraph("Install the Android/iOS app or use the web app described in the ", "official ntfy getting-started guide", LINKS["ntfy_docs"], ".")
        self.add_body("For the public HTTPS service, use a long random topic and an Apprise URL such as:")
        self.add_code_block(["ntfys://ntfy.sh/pdw-training-7fc2e91a-example"])
        self.add_body(
            "Public ntfy.sh topics are open by default: anyone who knows the topic can subscribe. For sensitive use, reserve/authenticate the topic where available or operate a private ntfy server with access control."
        )
        self.add_link_paragraph("Current secure ntfys syntax and authentication forms are documented by ", "Apprise ntfy service documentation", LINKS["apprise_ntfy"], ".")
        self.add_heading("Test sequence", 2)
        for step in (
            "Enter the authenticated HTTPS Apprise API URL and the correct stateless or stateful destination arrangement.",
            "Choose Send test notification. A test does not enable automatic delivery or save unsaved credentials.",
            "Confirm the device receives PDW Apprise test.",
            "Enable Apprise and save.",
            "Replay or inject a synthetic record that matches a training filter; confirm one notification.",
            "Replay a synthetic unfiltered record; confirm no filtered-only notification.",
            "Review Apprise.log for sanitized result categories if delivery fails.",
        ):
            self.add_step(step)

    def chapter_transfer(self) -> None:
        self.add_chapter("10. Secure file transfer", "secure-file-transfer")
        self.add_figure(
            "13-secure-file-transfer.png",
            "SFTP upload profile using fictional host and files.",
            "Continuous File Transfer dialog set to SFTP with fictional example host, username, password, remote folder, SHA256 host key, three published files and a 60-second interval.",
            width=6.2,
        )
        self.add_heading("Protocol choices", 2)
        self.add_table(
            ["Protocol", "Usual port", "Protection", "Use"],
            [
                ("FTP", "21", "None", "Legacy only. Username, password and files travel in plain text."),
                ("FTPS explicit", "21", "TLS upgrade", "Choose only when the host specifies explicit FTPS."),
                ("FTPS implicit", "990", "TLS from connection start", "Choose only when the host specifies implicit FTPS."),
                ("SFTP", "22", "SSH", "Preferred when the host provides password SFTP and a verified SHA-256 host key."),
            ],
            [1800, 1200, 2400, 3960],
            font_size=8.2,
        )
        self.add_heading("Get a hosting account", 2)
        self.add_body(
            "Choose a provider that explicitly supports password-authenticated SFTP or FTPS, supplies a public web folder when you are publishing a site, and can provide the server's verified SSH host-key fingerprint. PDW does not create remote folders."
        )
        self.add_link_paragraph("One example is a conventional web-hosting account such as ", "DreamHost web hosting", LINKS["dreamhost"], ". This is an example, not an endorsement; compare privacy, location, support and cost for your use.")
        self.add_link_paragraph("DreamHost's current SFTP credential guide shows where to obtain hostname, username, password and port 22: ", "FTP overview and credentials", LINKS["dreamhost_sftp"], ". Obtain the SHA-256 host key separately from the provider through a trusted channel.")
        self.add_heading("Configure an upload", 2)
        for step in (
            "Open Settings > Connections and automation > File transfer.",
            "Choose the exact protocol named by the provider. FTPS and SFTP are different protocols.",
            "Enter only the server hostname, then confirm the port, username and provider-specific remote folder.",
            "For SFTP, paste the provider's SHA256: host-key fingerprint. Do not copy it from an unexpected first-connection warning.",
            "Use Add files to select outputs. PDW rejects duplicate filenames because both would target the same remote name.",
            "Set an interval from 10 to 86400 seconds and select Enable automatic uploads.",
            "Choose Upload now for an immediate background transfer, then review the status and FileTransfer.log.",
        ):
            self.add_step(step)
        self.add_heading("Security behaviour", 2)
        self.add_bullet("The hosting password is a generic Windows Credential Manager entry scoped to the PDW installation; it is not written to PDW.INI or FileTransfer.log.")
        self.add_bullet("FTPS requires TLS 1.2 or later and validates certificate trust and hostname. There is no bypass switch.")
        self.add_bullet("SFTP authenticates only after the exact SHA-256 SSH host key matches.")
        self.add_bullet("Uploads use a temporary local snapshot so a changing logfile is transferred as a consistent file.")
        self.add_bullet("One background upload runs at a time; the next interval begins after it finishes.")
        self.add_bullet("The remote directory must exist and file names are retained.")
        self.add_callout(
            "Host-key rotation",
            "If a provider changes its SSH host key, stop transfers and verify the new fingerprint with the provider through a separate trusted channel before updating PDW.",
            kind="warning",
        )

    def chapter_publishing(self) -> None:
        self.add_chapter("11. Publish to a website or webhook", "publish-to-a-website-or-webhook")
        self.add_figure(
            "14-publishing-and-web.png",
            "Publishing profile with conservative privacy settings and fictional webhook values.",
            "Publishing and Web Integration dialog enabled for filtered messages, masked pager address, omitted message text, static files and an HTTPS Pipedream-style webhook with masked bearer and HMAC secrets.",
            width=6.25,
        )
        self.add_callout(
            "Acknowledgement required",
            "Publishing remains off until the operator acknowledges responsibility for checking permissions, privacy obligations, radio laws and publication laws in their jurisdiction. The acknowledgement is not legal advice or permission.",
            kind="danger",
        )
        self.add_heading("Start with a privacy-minimised profile", 2)
        self.add_bullet("Filtered messages only.")
        self.add_bullet("Mask pager address except the last three digits.")
        self.add_bullet("Leave Include decoded message text off.")
        self.add_bullet("Use a neutral public source alias that does not expose a site, person or receiver location.")
        self.add_bullet("Publish first to a private test destination using synthetic records.")
        self.add_heading("Static files", 2)
        self.add_table(
            ["File", "Purpose"],
            [
                ("messages.json", "Most recent events as a JSON feed."),
                ("messages.jsonl", "One JSON event per line for scripts or imports."),
                ("messages.rss", "RSS 2.0 feed."),
                ("messages.atom", "Atom feed."),
                ("index.html", "Responsive standalone browser table."),
            ],
            [2500, 6860],
        )
        self.add_body(
            "Select a local folder, then point File transfer at the generated files to place them on an SFTP/FTPS web host. A local web server can also serve the folder directly when properly secured."
        )
        self.add_heading("Generic HTTPS webhook", 2)
        self.add_body(
            "PDW sends one JSON object per event and refuses non-HTTPS URLs. Optional bearer and HMAC secrets are stored in Windows Credential Manager. An Idempotency-Key identifies the event. The receiving service must actually validate the bearer token or X-PDW-Signature; entering a secret in PDW alone does not secure an endpoint."
        )
        self.add_heading("Pipedream synthetic test example", 3)
        self.add_link_paragraph("Create a free account from the ", "Pipedream getting-started page", LINKS["pipedream"], ", then create a workflow with an HTTP / Webhook trigger.")
        for step in (
            "Copy the generated HTTPS endpoint into PDW's URL field.",
            "Keep message text off and use only synthetic records during initial testing.",
            "If using a bearer token or HMAC secret, add workflow logic that rejects missing or invalid headers before any downstream step.",
            "Choose Test webhook. PDW sends a configuration-only test object without decoded pager text.",
            "Inspect the event, then delete synthetic test events from the service if they are no longer required.",
            "Enable live delivery only after reviewing the provider's storage, retention, region, access controls and account security.",
        ):
            self.add_step(step)
        self.add_link_paragraph("Pipedream explains generated HTTP endpoints in its ", "HTTP trigger documentation", LINKS["pipedream_http"], ".")
        self.add_callout(
            "Third-party storage",
            "A webhook provider may retain request bodies, headers and event history. Do not send live decoded content until the service and workflow have been approved for that data.",
            kind="warning",
        )
        self.add_heading("Queue and failures", 2)
        self.add_bullet("Delivery is asynchronous and does not wait in the capture/decoder thread.")
        self.add_bullet("Pending events are stored under PublishQueue. Pause retains queued work.")
        self.add_bullet("Transient failures retry up to five times with exponential backoff.")
        self.add_bullet("Repeated failures move under PublishQueue\\DeadLetter for operator review.")
        self.add_bullet("The minimum interval provides simple rate limiting.")
        self.add_bullet("Privacy transforms change only the published copy; the display, logs, email and Apprise source record are not rewritten.")

    def chapter_operations(self) -> None:
        self.add_chapter("12. Daily operation and troubleshooting", "daily-operation-and-troubleshooting")
        self.add_heading("Start-of-shift checklist", 2)
        for item in (
            "Confirm the instance name, decoder mode and intended signal source.",
            "Confirm receiver/test status and the signal indicator without exposing live content in screenshots.",
            "Check date, time and available disk space.",
            "Confirm the correct log and filter files are open.",
            "Review Apprise, transfer and publishing status only if those channels are approved and enabled.",
            "Check PublishQueue\\DeadLetter and recent sanitized logs for unresolved failures.",
            "Run a synthetic filtered/unfiltered test after material configuration changes.",
        ):
            self.add_bullet(item)
        self.add_heading("Troubleshooting", 2)
        self.add_table(
            ["Symptom", "Checks and corrective action"],
            [
                ("No signal movement", "Confirm source, cable/device, Windows input permission, selected soundcard/COM port, receiver activity and volume/gain. Test the RTL source where applicable."),
                ("Activity but no messages", "Confirm mode, frequency, PPM, bandwidth and expected rate. Reduce clipping, check antenna and try known-good permitted traffic."),
                ("Poor local-audio decode", "Try the supplied audio presets and level changes before Custom. Record a short authorised sample and compare metrics."),
                ("Direct RTL-SDR not listed", "Close other SDR software, reconnect the device, confirm WinUSB on the exact interface, select the bundled pack and verify a compatible 32-bit DLL."),
                ("RTL-TCP test fails", "Confirm the host resolves, port is listening, firewall permits it and the server returns a valid RTL0 header. Use a trusted network."),
                ("Replay will not start", "Stop diagnostic recording, confirm supported WAV/SigMF encoding and path permissions, then retry."),
                ("Expected filter does not match", "Check type, address length/wildcards, function number, all & fragments, exact-match setting and Reject/Monitor Only state."),
                ("Email test fails", "Verify provider host, implicit SSL port, SMTP credentials, verified From address and network access. Use the relay's SMTP user password."),
                ("No Apprise notification", "Confirm the record reaches the lower filtered pane, Apprise is enabled, HTTPS/Basic auth works, and stateless destinations are not mixed with /notify/{key}."),
                ("SFTP upload fails", "Confirm host, port 22, password, existing remote folder and exact provider-verified SHA-256 host key."),
                ("FTPS certificate error", "Correct hostname and system trust/time. Do not bypass certificate validation."),
                ("Publishing test disabled/fails", "Enable a static or webhook output, acknowledge jurisdiction responsibility, use HTTPS and check pause/rate controls."),
                ("Queue grows", "Pause new exposure if needed, correct the destination, review pending/dead-letter data and re-enable only after a synthetic test."),
                ("Second PDW instance will not start", "Use a separate complete folder for each concurrent instance and give it a unique WindowTitle."),
            ],
            [2700, 6660],
            font_size=7.8,
        )
        self.add_heading("Diagnostic logs", 2)
        self.add_table(
            ["Log / folder", "Use", "Sensitive content"],
            [
                ("FileTransfer.log", "Transfer result and timing review.", "Designed to exclude the password; uploaded filenames and operational details still need protection."),
                ("Apprise.log", "Timestamp, generated event identifier and sanitized result category.", "Excludes decoded content, endpoint, credentials, destination URL and response body."),
                ("Published", "Inspect generated static outputs before upload.", "May contain published addresses, labels and message text according to the profile."),
                ("PublishQueue", "Pending webhook events.", "Contains event payloads; protect it as decoded data."),
                ("DeadLetter", "Repeated delivery failures.", "Contains failed event payloads; review and delete under policy."),
            ],
            [2300, 3100, 3960],
            font_size=8.0,
        )
        self.add_heading("Safe acceptance test", 2)
        for step in (
            "Copy the PDW folder to an isolated training location.",
            "Use a synthetic or explicitly licensed recording with one known matching and one known non-matching event.",
            "Confirm display fields, filter labels, hit counters and local logs.",
            "Use disposable third-party test destinations and no real tokens in screenshots or issue reports.",
            "Confirm filtered-only outputs send the matching record once and do not send the non-matching record.",
            "Deliberately test a wrong SFTP host key and an invalid TLS endpoint; both should fail closed.",
            "Remove disposable credentials and test data after acceptance.",
        ):
            self.add_step(step)
        self.add_heading("Backup and restore", 2)
        self.add_bullet("Close PDW before backing up settings.")
        self.add_bullet("Back up PDW.INI, filters.ini, required WAV/database files and receiver packages you intentionally imported.")
        self.add_bullet("Windows Credential Manager secrets are not present in the folder backup; re-enter them on the restored machine.")
        self.add_bullet("Do not copy private logs, recordings, Published or PublishQueue folders unless the backup is authorised and encrypted.")
        self.add_bullet("Restore into a separate folder first, verify settings, then retire the old copy only after acceptance.")

    def chapter_reference(self) -> None:
        self.add_chapter("13. Reference", "reference")
        self.add_heading("Menu reference", 2)
        self.add_table(
            ["Menu", "Commands"],
            [
                ("File", "Log file; Save copied data; Print copied data; Exit"),
                ("Edit", "Copy Selection; Copy monitor window; Copy filtered window; Clear monitor"),
                ("Monitor", "Input setup; Volume mixer; Decoding mode; Statistics; Pause or resume"),
                ("Filters", "Manage filters; Filter options; Reload filters; Reset all hit counters; Write filter file; Run filter command file"),
                ("Settings", "Settings hub; Appearance; Decoder and input; Connections and automation; Character set"),
                ("Help", "User guide; Diagnostics; About"),
            ],
            [1800, 7560],
            font_size=8.4,
        )
        self.add_heading("Keyboard shortcuts", 2)
        self.add_table(
            ["Shortcut", "Command", "Shortcut", "Command"],
            [
                ("Alt+L", "Log file", "Ctrl+C", "Copy selection"),
                ("Ctrl+U", "Copy monitor", "Ctrl+L", "Copy filtered"),
                ("Ctrl+D", "Clear display", "Ctrl+F", "Manage filters"),
                ("Ctrl+R", "Reload filters", "Ctrl+,", "Settings hub"),
                ("Ctrl+G", "General behaviour", "Ctrl+O", "Decoder options"),
                ("Ctrl+S", "Input setup", "Ctrl+M", "Email"),
                ("Ctrl+T", "System tray", "Alt+V", "Volume mixer"),
                ("Alt+S", "Statistics", "Alt+F", "Message font"),
                ("Ctrl+P", "Print copied data", "Ctrl+X", "Exit"),
                ("F1", "User guide", "F11", "Switch pane sizes"),
                ("F12", "Diagnostics", "Alt+Shift+R", "Legacy recording"),
                ("Alt+Shift+P", "Legacy playback", "Alt+Shift+A", "Legacy auto-record"),
            ],
            [1500, 3180, 1500, 3180],
            font_size=8.2,
        )
        self.add_heading("Radio defaults and limits", 2)
        self.add_table(
            ["Setting", "Default / range"],
            [
                ("RTL-TCP host", "127.0.0.1"),
                ("RTL-TCP port", "1234"),
                ("Frequency", "148000000 Hz initial value - replace with an authorised target"),
                ("IQ sample rate", "1024000 samples/s"),
                ("Audio rate", "48000 samples/s"),
                ("NFM bandwidth", "12000 Hz default; 5000 to 25000 Hz"),
                ("PPM correction", "0 initial"),
                ("Automatic gain", "Enabled initial"),
                ("Recording safety limit", "25 million normalized samples in memory"),
                ("Calibration search", "1,000 threshold/centering/resync combinations"),
            ],
            [3100, 6260],
        )
        self.add_heading("Third-party service directory", 2)
        service_rows = [
            ("SMTP2GO", "Email relay account and SMTP user", "SMTP host mail.smtp2go.com; PDW example uses SSL port 465."),
            ("Apprise API", "Operator-managed push gateway", "Host separately behind HTTPS and Basic authentication."),
            ("Pushover", "Phone/desktop push destination", "Account User Key plus an application API token."),
            ("ntfy", "Phone/web push destination", "Public or self-hosted topic; use secure ntfys URL and access controls."),
            ("Pipedream", "Optional webhook workflow", "Synthetic testing or reviewed automation; validate bearer/HMAC before downstream actions."),
            ("SFTP web host", "Static website/file destination", "Must provide credentials, remote folder and verified host key."),
            ("Zadig", "Optional WinUSB driver utility", "No account. Use only on the exact RTL-SDR interface."),
        ]
        self.add_table(["Service", "Why needed", "Key setup output"], service_rows, [1800, 3000, 4560], font_size=8.0)
        self.add_heading("Official links", 2)
        official_links = [
            ("PDW project site", LINKS["project"]),
            ("PDW source repository", LINKS["github"]),
            ("Apprise API", LINKS["apprise_api"]),
            ("Apprise supported services", LINKS["apprise_services"]),
            ("Pushover signup", LINKS["pushover_signup"]),
            ("Pushover Message API", LINKS["pushover_api"]),
            ("ntfy getting started", LINKS["ntfy_docs"]),
            ("SMTP2GO signup", LINKS["smtp2go_signup"]),
            ("SMTP2GO SMTP settings", LINKS["smtp2go_settings"]),
            ("Pipedream HTTP triggers", LINKS["pipedream_http"]),
            ("DreamHost SFTP credentials example", LINKS["dreamhost_sftp"]),
            ("Zadig", LINKS["zadig"]),
            ("RTL-SDR Blog quick start", LINKS["rtlsdr"]),
        ]
        for label, url in official_links:
            self.add_link_paragraph("", label, url)
        self.add_heading("Glossary", 2)
        self.add_table(
            ["Term", "Meaning"],
            [
                ("Capcode / Riccode", "Address used by a paging receiver and filter."),
                ("Discriminator tap", "Receiver output taken before normal audio filtering/de-emphasis for a cleaner data waveform."),
                ("FTPS", "FTP protected by TLS. It is not SFTP."),
                ("SFTP", "SSH File Transfer Protocol."),
                ("Host key", "SSH server identity fingerprint verified before password authentication."),
                ("HMAC", "Message authentication code computed over the exact webhook body with a shared secret."),
                ("Idempotency key", "Stable event identifier a receiver can use to avoid processing the same event twice."),
                ("IQ samples", "In-phase and quadrature radio samples received from an SDR."),
                ("NFM", "Narrowband frequency modulation."),
                ("SigMF", "Signal Metadata Format recording represented here as rf32_le data plus metadata."),
                ("WASAPI", "Current Windows audio capture API used as a fallback beside retained WinMM."),
                ("Dead letter", "Repeatedly failed queued event retained for operator review."),
            ],
            [2200, 7160],
            font_size=8.2,
        )
        self.add_heading("Credits and licence", 2)
        self.add_body("PDW credits retained from the application:")
        self.add_bullet("Jason Petty - original PDW project.")
        self.add_bullet("Peter Hunt - PDW development from 2004 to 2010.")
        self.add_bullet("Open-source release contributors, including later build and compatibility work.")
        self.add_bullet("Kieran O'Rourke - PDW v4.1.0 beta modernization contribution credited by the current About dialog.")
        self.add_body(
            "The supplied application identifies the GNU General Public License, version 3 or later. Review the included License and THIRD_PARTY_NOTICES.md files for the complete terms, bundled component versions, sources and licences."
        )
        self.add_callout(
            "Manual maintenance",
            "This edition reflects PDW v4.1.0 Beta and third-party documentation checked in August 2026. Re-check service URLs, account requirements, limits and security guidance before a later release or production rollout.",
            kind="note",
        )

    def set_core_properties(self) -> None:
        properties = self.document.core_properties
        properties.title = "PDW 4.1 Beta User Manual"
        properties.subject = "Operator guide for PDW Paging Decoder for Windows"
        properties.author = "PDW Community Project"
        properties.keywords = "PDW, POCSAG, FLEX, ACARS, MOBITEX, ERMES, RTL-SDR, Apprise"
        properties.comments = "All examples are fictional training data."

    def save(self) -> None:
        self.output.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(self.output)


class V551ManualBuilder(ManualBuilder):
    """Current operator manual for PDW v5.5.1 Public Beta 2."""

    def set_running_header_footer(self, section) -> None:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run("PDW v5.5.1 2026 Release User Manual")
        set_run_font(run, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_before = Pt(0)
        run = paragraph.add_run("PDW v5.5.1  |  Page ")
        set_run_font(run, size=8.5, color=MUTED)
        add_page_field(paragraph)

    def add_steps(self, steps: Sequence[str]) -> None:
        num_id = add_numbering_definition(self.document, ordered=True)
        for text in steps:
            paragraph = self.document.add_paragraph()
            apply_num(paragraph, num_id)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.line_spacing = 1.18
            set_run_font(paragraph.add_run(text))

    def add_toc(self) -> None:
        self.document.add_page_break()
        heading = self.document.add_paragraph("Contents", style="Heading 1")
        add_bookmark(heading, "contents", self.bookmark_id)
        self.bookmark_id += 1
        self.add_body(
            "Use the linked chapter names below or Word's Navigation Pane to move through the manual. Page numbers are intentionally omitted because they vary between Word/PDF renderers."
        )
        for index, (title, anchor) in enumerate(CHAPTERS, start=1):
            paragraph = self.document.add_paragraph(style="TOC Entry")
            add_internal_hyperlink(paragraph, f"{index}. {title}", anchor)

    def build(self) -> None:
        self.build_cover()
        self.build_front_matter()
        self.add_toc()
        self.chapter_read_first()
        self.chapter_install()
        self.chapter_interface()
        self.chapter_signal_source()
        self.chapter_decode()
        self.chapter_messages()
        self.chapter_filters()
        self.chapter_logs_display()
        self.chapter_notifications()
        self.chapter_transfer()
        self.chapter_publishing()
        self.chapter_data_outputs()
        self.chapter_operations()
        self.chapter_reference()
        self.set_core_properties()

    def build_cover(self) -> None:
        self.document.add_paragraph().paragraph_format.space_after = Pt(10)
        self.add_title_line("OPERATOR GUIDE", size=10, color=CYAN, after=13)
        self.add_title_line("PDW v5.5.1", size=31, color=NAVY, after=2)
        self.add_title_line("Paging Decoder for Windows", size=16, color=DARK_BLUE, bold=False, after=7)
        self.add_title_line("2026 Release - Public Beta 2", size=12, color=GOLD, after=12)
        self.add_figure(
            "01-main-window-v5.png",
            "Content-free PDW v5 command bar and live-input layout.",
            "Dark PDW main window with File, Monitor, Filters, Outputs, View and Help menus; Source, Pause, Record, Filters, Clear and Settings command buttons; a live input meter; monitored and filtered panes; and a status bar. No decoded traffic is shown.",
            width=6.45,
        )
        self.add_title_line("Edition 2.0  |  11 August 2026  |  Covers PDW v5.5.1 Public Beta 2", size=9.3, color=MUTED, after=4)
        self.add_title_line("Native x64 and Win32 applications for supported Windows 10/11 and Windows Server systems", size=8.8, color=MUTED, bold=False, after=3)
        self.add_title_line("All addresses, messages, accounts, endpoints and credentials in this guide are synthetic or masked.", size=8.4, color=RED, bold=True, after=0)

    def build_front_matter(self) -> None:
        self.document.add_page_break()
        self.add_heading("How to use this guide", 1)
        self.add_callout(
            "Training data",
            "Use only synthetic, redacted or explicitly licensed examples for setup, screenshots and acceptance. Never place private pager traffic, recordings, credentials, Windows endpoint IDs or operator settings in a public report.",
            kind="warning",
        )
        self.add_body(
            "This manual replaces the 2010 PDW v3.1 guide and the interim v4.1 draft. It keeps the established protocol, decoder, filter and logging concepts, but follows the v5.5.1 installer, dual-architecture packages, 2026 command bar, Settings Center, Capcode Directory, message history, secure outputs and named local-audio profile."
        )
        self.add_heading("Feature status used in this manual", 2)
        self.add_table(
            ["Status", "Meaning"],
            [
                ("Public Beta 2", "Included in the published v5.5.1 prerelease. It remains unsigned and is not a signed stable release."),
                ("Preserved legacy", "Established decoder, serial, slicer, WinMM, filter or display behaviour intentionally retained for compatibility."),
                ("Optional / off", "Included but disabled until an operator deliberately enables and configures it."),
                ("Hardware-unverified", "Automated checks pass, but representative physical-device acceptance remains incomplete."),
                ("Development preview", "Visible in the active 11 August 2026 source worktree, but not part of the published Public Beta 2 contract."),
            ],
            [2200, 7160],
            font_size=8.5,
        )
        self.add_heading("Ten-minute first run", 2)
        self.add_steps((
            "Download the v5.5.1 Public Beta 2 Setup only from the maintained GitHub release and verify its published SHA-256 checksum.",
            "Run Setup, read the unknown-publisher warning, and choose x64 for normal current Windows use or Win32 only when a legacy 32-bit receiver or integration requires it.",
            "Choose Standard PDW settings. Select the Adelaide FLEX profile only for the documented SDR# and VB-Audio Cable workflow.",
            "Start PDW, open Settings, and choose the intended source under Signal & radio. Use only transmissions you are authorised to receive.",
            "Select POCSAG/FLEX, ACARS, MOBITEX or ERMES and confirm the live meter responds without persistent clipping.",
            "Open Filters > Capcode Directory and create one synthetic training rule. Keep external output routing off for the first test.",
            "If local history is required, enable it under View > Message history and decide separately whether message text may be retained.",
            "Run configuration-only tests for any approved email, notification, transfer, publishing or data-output destination before enabling live delivery.",
            "Create an encrypted Settings > General > Backup / Restore file after configuration and protect its password separately.",
        ))

    def chapter_read_first(self) -> None:
        self.add_chapter("1. Read this first", "read-this-first")
        self.add_heading("What PDW does", 2)
        self.add_body(
            "PDW is a native Windows application that receives, decodes, displays, filters, logs and optionally routes supported digital radio messages. Current development is deliberately additive: modern Windows input, user-interface and output features sit beside established decoder and hardware paths rather than replacing them."
        )
        self.add_table(
            ["Protocol", "Supported role", "Compatibility boundary"],
            [
                ("POCSAG", "512, 1200 and 2400 baud", "Established parser and serial/audio routes retained."),
                ("FLEX", "1600, 3200 and 6400 baud", "Established parser retained; optional fragment assistance is separately gated."),
                ("ERMES", "6250 baud", "Established decoder route retained."),
                ("ACARS", "Aircraft data messages", "Parity and label/database options remain operator controlled."),
                ("MOBITEX", "8000 baud in normal use", "Established base/mobile and message controls retained."),
            ],
            [1500, 2450, 5410],
            font_size=8.3,
        )
        self.add_heading("What Public Beta 2 includes", 2)
        self.add_table(
            ["Area", "Included capability", "Default / evidence boundary"],
            [
                ("Windows", "Guided Setup, portable packages, x64 and Win32, light/dark UI and modeless Settings Center.", "Unsigned beta; full DPI, keyboard, High Contrast and physical-radio matrices remain open."),
                ("Signal", "WinMM, exact-endpoint WASAPI, serial/slicer, RTL-TCP, RTL-SDR, WAV/SigMF record and replay.", "Named Adelaide profile is opt-in and hardware-unverified."),
                ("Local tools", "Capcode Directory, optional SQLite history, CSV export, local dashboard and isolated extra receiver workers.", "History, dashboard and extra channels are off until enabled."),
                ("Routing", "Email, Apprise, FTPS/SFTP, publishing, MQTT, SQLite, MySQL/ODBC, Telnet and Windows notifications.", "Every network/output path is disabled by default and must fail independently from decoding."),
                ("FLEX", "Optional bounded fragment reassembly adds a marked assembled copy while retaining original fragments.", "Disabled by default; recording-backed live acceptance remains open."),
            ],
            [1500, 4500, 3360],
            font_size=7.8,
        )
        self.add_callout(
            "Privacy and law",
            "PDW cannot determine whether reception, storage, forwarding or publication is lawful in your location. Confirm authority, radio rules, privacy obligations, retention and publication requirements before use. Treat capcodes, aliases and decoded content as potentially identifying information.",
            kind="danger",
        )
        self.add_callout(
            "Release status",
            "Public Beta 2 passed automated x64/Win32 build, test, package, installer and Defender gates. It is intentionally unsigned, and physical SDR, SDR#, VB-CABLE, legacy receiver and complete UI acceptance remain incomplete. Do not describe it as signed, stable or fully hardware-validated.",
            kind="warning",
        )

    def chapter_install(self) -> None:
        self.add_chapter("2. Install and first start", "install-and-first-start")
        self.add_heading("Supported Windows boundary", 2)
        self.add_body(
            "Setup enforces Windows build 10586 as the technical API floor. Normal use requires a Windows 10 or Windows 11 edition/build still receiving Microsoft security servicing, or a serviced Windows Server 2016-or-newer release. Windows 7, 8 and 8.1 are outside the current Visual Studio 2026 and Windows SQLite support boundary."
        )
        self.add_heading("Choose an architecture", 2)
        self.add_table(
            ["Package", "Choose it for", "Native dependency rule"],
            [
                ("x64", "Current 64-bit Windows, Windows audio, RTL-TCP and matching x64 receiver libraries.", "All loaded native DLLs must be x64."),
                ("Win32", "Bundled x86 RTL-SDR, historical slicer/receiver libraries or another known 32-bit integration.", "All loaded native DLLs must be x86/Win32."),
            ],
            [1500, 4900, 2960],
        )
        self.add_callout("Architecture", "A 32-bit DLL cannot load into x64 PDW, and a 64-bit DLL cannot load into Win32 PDW. Decoder, filter and configuration behaviour are shared between the builds.", kind="note")
        self.add_heading("Guided Setup", 2)
        self.add_steps((
            "Verify the downloaded Setup checksum against the value published with Public Beta 2.",
            "Run Setup and choose the destination folder. The default is under the current user's Local Programs folder.",
            "On 64-bit Windows, choose x64 or Win32 compatibility. Setup selects Win32 automatically on 32-bit Windows.",
            "For a new settings file, choose Standard PDW settings or the optional SDR# + VB-Audio Cable (Adelaide FLEX) profile.",
            "Choose Start Menu, optional Desktop and optional delayed Windows-startup shortcuts.",
            "Launch PDW and confirm the title/About identity and selected architecture before configuring hardware or outputs.",
        ))
        self.add_callout(
            "Unknown publisher",
            "Public Beta 2 is intentionally unsigned. Windows or SmartScreen may identify an unknown publisher. Download only from the maintained release, verify the checksum, and do not bypass the warning for an unverified copy.",
            kind="warning",
        )
        self.add_heading("Upgrade and uninstall", 2)
        self.add_bullet("Normal upgrades preserve PDW.INI, Capcode Directory/history database, receiver additions, WAV files, logs, recordings, queues and other operator-created data.")
        self.add_bullet("A pre-existing PDW.INI always wins; Setup does not silently apply a clean-install profile during an upgrade.")
        self.add_bullet("Architecture changes preserve the active RTL-SDR DLL to a verified recovery backup before replacing or removing the active slot.")
        self.add_bullet("Uninstall leaves configuration and operator data by default. Review and remove it manually only after an authorised backup.")
        self.add_heading("Portable packages", 2)
        self.add_body("The portable x64 and Win32 packages remain supported. Extract the complete package to a writable folder and keep the executable, PDW.INI, Capcode Directory/history database, receivers, WAV files and logs together. Do not run PDW from inside a ZIP file or mix files from different releases.")
        self.add_heading("Important local data", 2)
        self.add_table(
            ["Path / item", "Purpose and handling"],
            [
                ("PDW.INI", "Primary saved configuration beside PDW. Protect it as operator configuration."),
                ("pdw-history.sqlite3", "Capcode Directory and, when enabled, message history. The directory remains usable with history capture off."),
                ("filters.ini.migrated", "Recoverable backup created after a successful one-time legacy filter migration."),
                ("Wavfiles / Receivers", "Optional alert audio and receiver packages. Native receiver DLLs must match PDW architecture."),
                ("Logfiles / Published / queues", "May contain decoded or transformed content. Apply access, retention and disposal controls."),
                (".pdwbackup", "Password-encrypted configuration export that can include supported saved credentials."),
            ],
            [2500, 6860],
            font_size=8.2,
        )

    def chapter_interface(self) -> None:
        self.add_chapter("3. Interface tour", "interface-tour")
        self.add_figure(
            "01-main-window-v5.png",
            "PDW v5 command bar, panes, live meter and status bar (content-free native capture).",
            "Dark PDW main window with six menus, six labelled command buttons, a green live-input meter, empty monitor and filtered panes, and a four-part status bar.",
            width=6.5,
        )
        self.add_heading("Main window", 2)
        self.add_table(
            ["Area", "Use"],
            [
                ("Menus", "File, Monitor, Filters, Outputs, View and Help organise the complete command set."),
                ("Command bar", "Source, Pause, Record, Filters, Clear and Settings provide the primary daily actions."),
                ("Live input meter", "Shows source state, source name, level/quality and real sample activity. It is not calibrated RF power."),
                ("Monitor pane", "Upper pane for displayed decoded messages."),
                ("Filtered pane", "Lower pane for rules whose Filter action is enabled and that are not rejected or Monitor Only."),
                ("Status bar", "Reports source, decoder, capture state and output-health summary without displaying decoded text."),
            ],
            [2200, 7160],
            font_size=8.5,
        )
        self.add_heading("The six menus", 2)
        self.add_table(
            ["Menu", "Main tasks"],
            [
                ("File", "Log, save/print copied data, configuration backup/restore and exit."),
                ("Monitor", "Pause, decoder mode, input setup, radio/replay, statistics and volume."),
                ("Filters", "Capcode Directory, filter controls, counters and legacy compatibility actions."),
                ("Outputs", "Email, Apprise, transfer, publishing, data outputs, Delivery Health, dashboard and guarded channels."),
                ("View", "Message history, display/column options, appearance and window tools."),
                ("Help", "User guide, diagnostics, About, licence and project links."),
            ],
            [1750, 7610],
            font_size=8.4,
        )
        self.add_heading("Settings Center", 2)
        self.add_figure(
            "02-settings-center-v5.png",
            "Modeless Settings Center navigation used by the v5 family.",
            "Dark Settings Center with General, Appearance, Display, Decoder, Signal and radio, Filters, Notifications, Data outputs, Health and diagnostics, and About me navigation entries.",
            width=6.0,
        )
        self.add_body("Settings is modeless: monitoring continues while it is open. Draft changes are retained while moving between pages, and retained legacy editors return to the same page when closed. Opening Settings again focuses the existing window instead of creating a second copy.")
        self.add_callout("Accessibility", "Use clear message fonts and only the columns you need. Confirm keyboard navigation, High Contrast and 125-200% scaling on the intended computer before relying on the beta in an accessibility-sensitive workflow.", kind="note")
        self.add_heading("Development preview: consolidated Settings", 2)
        self.add_figure(
            "03-settings-development-preview.png",
            "11 August 2026 development preview of consolidated Data outputs navigation.",
            "Dark PDW v5.5.1 Settings development screen with nine navigation pages and Email, Push and Windows notifications, File transfer, Publish to web and Data outputs cards on one Data outputs page.",
            width=6.0,
        )
        self.add_callout("Development preview", "The active source worktree consolidates duplicate Settings cards and groups email, push/Windows notifications, transfer, publishing and optional data outputs on one page. This screenshot is not a promise that the published Public Beta 2 menu count or wording has changed.", kind="warning")

    def chapter_signal_source(self) -> None:
        self.add_chapter("4. Choose a signal source", "choose-a-signal-source")
        self.add_heading("Source decision guide", 2)
        self.add_table(
            ["Source", "Use it when", "Important boundary"],
            [
                ("Local audio", "A receiver or virtual cable supplies audio.", "Identityless profiles try legacy WinMM first; saved named endpoints open exact-ID WASAPI and fail closed if missing."),
                ("Serial / slicer / RS232", "Known compatible legacy hardware is available.", "Win32 remains the compatibility choice for x86-only integrations; obsolete VXD packages are not shipped."),
                ("RTL-TCP", "A local or trusted-network rtl_tcp-compatible server supplies IQ.", "Architecture-neutral TCP path; enhanced IQ filtering is optional and disabled by default."),
                ("Direct RTL-SDR", "A compatible USB receiver is attached.", "Win32 includes the reviewed x86 receiver pack; x64 needs a matching trusted x64 library or RTL-TCP."),
                ("Replay", "A synthetic, redacted or licensed WAV/SigMF recording is available.", "Runs through normal decoder functions and restores the previous live source."),
            ],
            [1500, 3000, 4860],
            font_size=8.0,
        )
        self.add_heading("Local audio and serial", 2)
        self.add_figure(
            "03-legacy-input-setup.png",
            "Legacy Input Setup remains available beside modern source controls.",
            "PDW Interface Setup dialog with Serial Port and Soundcard groups, configuration preset, sample rate and device controls.",
            width=4.3,
        )
        self.add_steps((
            "Connect the intended receiver output, virtual cable or compatible serial interface before opening PDW.",
            "Open Settings > Signal & radio and choose the current source or retained legacy input editor.",
            "For audio, select the exact device and the closest supported configuration. Avoid persistent clipping.",
            "For serial, select the exact COM port, Slicer or RS232 mode and the required two- or four-level path.",
            "Select the matching protocol family and confirm valid synthetic/licensed traffic before saving automation settings.",
        ))
        self.add_heading("SDR# + VB-Audio Cable (Adelaide FLEX)", 2)
        self.add_body("The named profile is an explicit clean-install or operator-confirmed choice. SDR# sends unfiltered NFM audio to CABLE Input; PDW captures the matching CABLE Output recording endpoint. PDW stores the exact Windows endpoint identity and does not silently fall back to a microphone when that identity is missing or ambiguous.")
        self.add_bullet("SDR# and VB-Audio Cable are separate products; PDW does not install, license, tune or configure them.")
        self.add_bullet("The documented profile uses 148.8125 MHz, NFM, Filter Audio off, unmuted audio, FLEX 1600 and the audited decoder/slicer values.")
        self.add_bullet("Existing/portable users may preview Apply Adelaide FLEX. The confirmation defaults to No and creates a verified PDW.INI backup before an approved change.")
        self.add_bullet("The profile does not create, replace or delete Capcode Directory rules, message history, receivers or WAV files.")
        self.add_callout("Hardware-unverified", "Automated endpoint, fail-closed, profile and installer checks pass, but the exact physical SDR#/VB-CABLE workflow is not maintainer-validated for Public Beta 2. Describe it as configured or operator-reported until representative acceptance is recorded.", kind="warning")
        self.add_heading("RTL-TCP and direct RTL-SDR", 2)
        self.add_figure(
            "04-radio-and-replay.png",
            "RTL-TCP setup using fictional training values.",
            "Radio and Signal Sources dialog set to RTL-TCP with fictional host decoder-gateway.example.net, port 1234, frequency 148337500 Hz and diagnostic controls.",
            width=6.0,
        )
        self.add_body("RTL-TCP validates the RTL0 header, configures the tuner, demodulates NFM and reconnects after a drop. Direct RTL-SDR lists devices and checks library architecture/API compatibility. Add receiver imports only a trusted matching-bitness librtlsdr package. PDW never launches Zadig or changes a USB driver automatically.")
        self.add_heading("Record, replay and diagnose", 2)
        self.add_bullet("WAV records normalized mono 16-bit samples; SigMF records rf32_le plus sample-rate metadata.")
        self.add_bullet("Replay accepts supported PCM8, PCM16 or float32 WAV and real float32 SigMF, resets timing, then restores the prior source.")
        self.add_bullet("The diagnostics view includes waveform, spectrum, waterfall, quality history, level, noise, clipping, eye opening and FLEX phase/error totals.")
        self.add_bullet("Calibration tests 1,000 threshold/centering/resync combinations and offers an operator-approved suggestion; it is not protocol-level proof.")
        self.add_heading("Guarded multi-channel receivers", 2)
        self.add_body("Outputs > Guarded multi-channel receivers can run up to four additional rtl_tcp or direct RTL-SDR channels. Each uses an isolated PDW worker process because the established decoders contain process-global state. Local history must be enabled, and each endpoint/device must be unique.")
        self.add_callout("Isolation", "Workers write channel-labelled events to shared local history but disable publishing, email, Apprise, data outputs, FTP, Telnet, dashboard and legacy logs. This release does not split one wideband IQ stream into several tuned channels.", kind="note")

    def chapter_decode(self) -> None:
        self.add_chapter("5. Decode and improve a signal", "decode-and-improve-a-signal")
        self.add_heading("Choose the decoder family", 2)
        self.add_body("Use Monitor > Decoding mode or the Source/Mode route and select POCSAG/FLEX, ACARS, MOBITEX or ERMES. PDW monitors one protocol family at a time. Activity without valid messages often indicates the wrong family, tuning, bandwidth, level or source.")
        self.add_heading("Improve reception in the right order", 2)
        self.add_steps((
            "Confirm the source and protocol family.",
            "Confirm frequency, receiver mode, sample rate and expected baud/rate.",
            "Correct gross under-level or clipping and improve the antenna/signal path.",
            "For SDR, adjust PPM, gain/AGC and NFM bandwidth using the same authorised recording or signal.",
            "Compare the live meter, waveform, eye/phase information and corrected/uncorrectable totals.",
            "Only then change custom thresholds or apply a calibration suggestion, keeping a reversible backup.",
        ))
        self.add_heading("Optional RTL IQ conditioning", 2)
        self.add_body("Enhanced IQ filtering and resampling applies only to RTL input. It uses a 60 dB windowed-sinc channel filter and anti-alias resampling before the existing decoder boundary. It is disabled by default; with it off, the established one-pole/averaging path remains. Compare the same lawful diagnostic recording before retaining the option.")
        self.add_heading("Duplicate and message boundaries", 2)
        self.add_body("Duplicate blocking, reject rules and protocol-specific display rules remain ahead of external output delivery. Optional output adapters receive an immutable event copy and cannot change the decoder, display or original filter result.")
        self.add_heading("Public Beta 2 FLEX fragment assistance", 2)
        self.add_body("Public Beta 2 can optionally assemble a complete non-group FLEX alpha/secure fragment chain. The feature is disabled by default. Original fragments continue through the established path, and one additional compact message marked [Joined FLEX] is emitted only after a bounded valid chain completes. Missing, conflicting, timed-out or capacity-rejected observations do not create a guessed joined message. FLEX Group Mode remains on its legacy path.")
        self.add_callout("Live evidence", "Automated fragment-order, replay, timeout, capacity and identity tests do not replace recording-backed live acceptance. Use only synthetic, redacted or licensed recordings.", kind="warning")
        self.add_heading("Development preview: joined-only and Part X of Y", 2)
        self.add_figure(
            "05-screen-options-preview.png",
            "Development preview of the joined-only FLEX screen option.",
            "Dark Screen Options dialog with seven column selectors, FLEX Group Mode controls and a Wait for complete split FLEX alpha or secure message option.",
            width=4.6,
        )
        self.add_body("The 11 August 2026 worktree changes the enabled FLEX option to hold valid fragments until one joined event enters display, filtering, logging and routing. It also recognises bounded explicit text markers such as Part 1 of 2 and can join 2-32 parts for one visible identity. These changes are development preview, not Public Beta 2 behaviour, and remain subject to build, test, package and recording-backed acceptance.")

    def chapter_messages(self) -> None:
        self.add_chapter("6. Understand decoded messages", "understand-decoded-messages")
        self.add_callout("Synthetic examples", "The addresses, times, names and messages below are invented training data and do not identify a real person, aircraft or service.", kind="note")
        self.add_heading("Paging message fields", 2)
        self.add_table(
            ["Address", "Time", "Mode", "Type", "Bitrate", "Synthetic message"],
            [
                ("1234567", "14:17:25", "POCSAG-1", "ALPHA", "1200", "TEST ONLY - maintenance drill complete"),
                ("1765432", "14:19:08", "FLEX-A", "ALPHA", "1600", "TRAINING - airport operations standby"),
                ("1002048", "14:22:31", "FLEX-C", "NUMERIC", "3200", "5550102"),
            ],
            [1200, 1000, 1450, 1150, 1000, 3560],
            font_size=7.8,
        )
        self.add_body("Address remains the raw decoded capcode/Riccode. Mode identifies protocol and function/phase. Type distinguishes alpha, numeric, tone-only or other protocol-specific content. Bitrate records the detected/selected rate. A current Capcode Directory name may be displayed beside the raw address but never replaces it.")
        self.add_heading("ACARS and MOBITEX", 2)
        self.add_body("ACARS records may include aircraft address/registration, flight, label, block/ack and message text. MOBITEX records may include network, base/mobile identity and protocol-specific message type. Keep only the columns and databases needed for an authorised workflow.")
        self.add_heading("Capcode Directory", 2)
        self.add_body("The local SQLite archive stores operator-maintained names, agencies, colours, notes and filter/routing rules. Protocol-specific mappings win over all-protocol mappings. The raw capcode remains available for audit and export.")
        self.add_heading("Message history", 2)
        self.add_body("Open View > Message history. History capture is disabled by default, and storing message text is a separate disabled-by-default choice. Retention may be set from 1 to 3650 days. The Capcode Directory remains available when history capture is off, and purging message history does not delete directory entries.")
        self.add_bullet("The default archive is pdw-history.sqlite3 beside PDW unless another path is selected.")
        self.add_bullet("The viewer searches and pages through 200-row displays without implying that only the visible page exists.")
        self.add_bullet("Corrupt or unrelated SQLite databases are refused unchanged; PDW does not silently repair or convert them.")
        self.add_heading("Export message history to CSV", 2)
        self.add_body("Export CSV writes every stored row matching the current Search, Protocol and Filtered controls, not just the displayed page. The UTF-8 file uses current enabled directory aliases at export time and leaves Message blank when text was not retained.")
        self.add_code_block(["Received,Protocol,Capcode,Name,Agency,Type,Message,Filter"])
        self.add_callout("Exported data", "A CSV is an unencrypted copy outside SQLite retention and purge controls. Formula-leading values are protected, but the file still requires authorised storage, sharing and deletion.", kind="warning")
        self.add_heading("Local live dashboard", 2)
        self.add_body("Outputs > Local live dashboard provides a read-only browser view and JSON API. It is disabled by default, requires message history, binds only to 127.0.0.1 and accepts only localhost/127.0.0.1 Host headers. It cannot accept remote connections or change PDW.")

    def chapter_filters(self) -> None:
        self.add_chapter("7. Build filters and route selected messages", "build-filters-and-alerts")
        self.add_figure(
            "06-capcode-routing-v5.png",
            "Capcode Directory rule editor with explicit output destinations (content-free native capture).",
            "Dark Capcode Directory and Filters window with empty directory, protocol, filter type, capcode, display name, agency, keyword, Filter, Reject, Monitor Only, output destination, CSV and general controls. No live capcodes or messages are shown.",
            width=6.45,
        )
        self.add_heading("One directory, three separate decisions", 2)
        self.add_table(
            ["Decision", "Control", "Effect"],
            [
                ("Display", "Display name, agency placement, colour", "Adds operator-maintained context while retaining the raw capcode."),
                ("Pane", "Filter, Monitor Only or Reject", "Sends a match to the lower pane, keeps it in the upper pane, or suppresses it."),
                ("Delivery", "Send to enabled outputs plus named destinations", "Routes only to destinations that are also enabled/configured in Settings."),
            ],
            [1700, 3000, 4660],
            font_size=8.4,
        )
        self.add_heading("Create a rule", 2)
        self.add_steps((
            "Open Filters > Capcode Directory and choose New.",
            "Choose one protocol or Any protocol, then enter the exact capcode/address.",
            "Add a synthetic display name, optional agency/service, colour and notes.",
            "Choose capcode-only, one required keyword, two-to-ten required keywords joined with +, or exact whole-message matching.",
            "Choose Filter, Monitor Only or Reject. Filter and Monitor Only are mutually exclusive.",
            "Leave Send to enabled outputs off until each intended destination is separately approved, configured and tested.",
            "Save or update. The live in-memory filter reloads immediately without replacing decoder logic.",
        ))
        self.add_heading("Matching precedence", 2)
        self.add_body("For the same capcode, PDW evaluates the most specific message condition first: exact whole-message, multi-keyword, single-keyword, then capcode-only. Keyword matching ignores case. A + expression requires every listed keyword in any order. Legacy & expressions remain supported for compatibility.")
        self.add_heading("Per-rule outputs", 2)
        self.add_body("A rule may select Email, Apprise, Publishing, MQTT, SQLite, MySQL/ODBC, Telnet and Windows notifications. Selecting a destination on the rule does not enable it. The destination and any required output-group/privacy acknowledgement must also be enabled in Settings. This two-step gate prevents a directory row from silently activating external delivery.")
        self.add_heading("CSV import/export and legacy migration", 2)
        self.add_bullet("Directory Import and Export use UTF-8 CSV. Invalid rows are rejected and reported.")
        self.add_bullet("Fresh packages do not include filters.ini. If an older readable file exists, PDW performs a transactional one-time merge into the directory.")
        self.add_bullet("After success, the original file is renamed to a unique .migrated recovery backup; a failed import leaves it unchanged.")
        self.add_bullet("Retire any scheduled legacy generator after migration and use the directory CSV controls for future bulk maintenance.")

    def chapter_logs_display(self) -> None:
        self.add_chapter("8. Log, display, copy and review", "log-copy-display-and-review")
        self.add_heading("Display and columns", 2)
        self.add_body("View/display settings control the seven message columns, colours, fonts, scrollback and pane arrangement. Keep the raw address, time and protocol/type information needed for interpretation, but remove unused columns to improve readability. Long v5 messages can use the wider layout without changing decoder content.")
        self.add_figure(
            "05-screen-options-preview.png",
            "Current development Screen Options layout; the seven column selectors also exist in the v5 family.",
            "Dark Screen Options dialog with seven column selectors and FLEX display controls.",
            width=4.6,
        )
        self.add_heading("Copy, save and print", 2)
        self.add_bullet("Select decoded rows or text, then use Edit/Copy or the appropriate menu command.")
        self.add_bullet("Save copied data creates a separate operator-controlled file; it is not covered by message-history retention.")
        self.add_bullet("Print only to an approved device or secured PDF location and dispose of output under policy.")
        self.add_heading("Main logfile", 2)
        self.add_body("Logging is operator-controlled. Choose a protected location and confirm whether monitor and filtered content, labels and line layout meet the authorised purpose. A logfile may contain private decoded content even when external publishing has message omission enabled.")
        self.add_heading("Separate filter CSV files", 2)
        self.add_body("A directory rule can write the same matching row to up to three selected CSV paths using the shared subset of Capcode, Time, Date, Mode, Type, Bitrate and Message. These local files are independent from Message History CSV export and require their own access/retention controls.")
        self.add_heading("Data boundary", 2)
        self.add_table(
            ["Store/output", "Contains", "Retention boundary"],
            [
                ("Main logfile", "Displayed/logged decoded records according to local settings.", "Operator-controlled file."),
                ("Message history", "Optional bounded SQLite events; text optional.", "1-3650 day configured retention and purge."),
                ("History CSV", "All stored rows matching current viewer filters.", "Unencrypted export outside archive retention."),
                ("Published / queues", "Transformed publishing copies and pending/dead-letter events.", "Publishing profile and manual queue review."),
                ("Delivery Health", "Destination names, counters, timestamps and sanitized outcomes.", "Never stores capcodes or decoded message text."),
            ],
            [2100, 3900, 3360],
            font_size=8.0,
        )

    def chapter_notifications(self) -> None:
        self.add_chapter("9. Email and notifications", "email-and-push-notifications")
        self.add_heading("Two independent message gates", 2)
        self.add_body("A Capcode Directory rule first decides whether a message is filtered/displayed and whether it selects an output. The email, Apprise or Windows-notification destination must then be enabled and configured in Settings. A rule never activates a destination by itself.")
        self.add_heading("SMTP email", 2)
        self.add_figure(
            "11-email-settings.png",
            "SMTP configuration using fictional, masked training values.",
            "PDW email settings dialog with fictional server and masked password fields.",
            width=5.6,
        )
        self.add_bullet("Use a current authenticated SMTP relay account and the provider's required TLS/port settings.")
        self.add_bullet("Test with a synthetic message and a non-production recipient before enabling matching-message delivery.")
        self.add_bullet("Do not place mail passwords or decoded traffic in screenshots, logs or issue reports.")
        self.add_heading("Apprise filtered-message push", 2)
        self.add_figure(
            "12-apprise-notifications.png",
            "Apprise notification settings with masked fictional values.",
            "Apprise dialog with Enable Apprise, masked API URL, username, password and destination fields; decoded text is not selected.",
            width=5.7,
        )
        self.add_body("PDW connects to an operator-managed HTTPS Apprise API. Public Beta 2 sends only messages selected by the established filtered/output routing boundary. Monitor Only, rejected, duplicate and unselected messages are not pushed. Delivery uses a bounded background queue so an unavailable phone service does not block decoding.")
        self.add_bullet("Use HTTPS and authentication; do not expose an unauthenticated Apprise API to an untrusted network.")
        self.add_bullet("Keep decoded text off unless lock-screen exposure has been specifically approved.")
        self.add_bullet("Use Send test notification before Enable Apprise. Testing must use synthetic content.")
        self.add_heading("Native Windows notifications", 2)
        self.add_body("Windows notifications are a separate optional data-output destination. They load/register only when enabled or tested. Message text has its own opt-in because notifications may appear on the lock screen. Turning the adapter off leaves PDW's legacy tray behaviour unchanged.")
        self.add_callout("Secrets", "Supported destination credentials are held through Windows Credential Manager and the encrypted backup workflow. Never rely on an ordinary folder copy as a safe secret backup, and never publish an exported configuration file or its password.", kind="warning")

    def chapter_transfer(self) -> None:
        self.add_chapter("10. Secure file transfer", "secure-file-transfer")
        self.add_figure(
            "13-secure-file-transfer.png",
            "Secure file-transfer settings using fictional training values.",
            "PDW file-transfer dialog with fictional host, username, masked password, protocol and remote folder values.",
            width=5.8,
        )
        self.add_heading("Supported protocols", 2)
        self.add_table(
            ["Protocol", "Protection", "Required verification"],
            [
                ("FTP", "No transport encryption", "Use only inside an explicitly approved protected environment."),
                ("Explicit FTPS", "TLS upgrade on the FTP control connection", "Windows-trusted certificate and hostname validation."),
                ("Implicit FTPS", "TLS from connection start", "Windows-trusted certificate and hostname validation."),
                ("SFTP", "SSH file transfer with password authentication", "Exact operator/provider-verified SHA-256 host-key fingerprint."),
            ],
            [1700, 3900, 3760],
            font_size=8.2,
        )
        self.add_heading("Configure an upload", 2)
        self.add_steps((
            "Create and approve the remote account/folder outside PDW.",
            "Choose the secure protocol, host, port, username, remote folder and local source file.",
            "For SFTP, obtain the exact SHA-256 host-key fingerprint through a trusted channel. For FTPS, use the correct certificate hostname.",
            "Enter the password through PDW's credential control and run the configuration test.",
            "Confirm the synthetic/local test file arrives in the intended folder before enabling continuous transfer.",
            "Set a reasonable interval and review sanitized FileTransfer.log outcomes.",
        ))
        self.add_callout("Fail closed", "Do not bypass certificate or host-key validation to make a test pass. A changed SFTP host key or certificate error must be investigated through an independent trusted channel.", kind="danger")

    def chapter_publishing(self) -> None:
        self.add_chapter("11. Publish to a website or webhook", "publish-to-a-website-or-webhook")
        self.add_figure(
            "14-publishing-and-web.png",
            "Publishing settings using fictional training values.",
            "PDW publishing dialog with static feed and HTTPS webhook controls, privacy options and masked secret fields.",
            width=5.9,
        )
        self.add_heading("Start disabled and minimise data", 2)
        self.add_body("Publishing is disabled by default and requires an operator acknowledgement of permissions, privacy, radio and publication responsibility. Start with filtered-only delivery, address masking and message omission. Those transforms affect only the published copy; they do not rewrite the local display, logs, email or Apprise source event.")
        self.add_heading("Static publishing", 2)
        self.add_body("PDW can generate JSON, JSONL, RSS, Atom and responsive HTML. Generate to a private local folder first, inspect only synthetic output, then use the separately configured transfer path if remote upload is approved.")
        self.add_heading("HTTPS webhook", 2)
        self.add_bullet("Use an HTTPS endpoint with reviewed retention, region, access and account controls.")
        self.add_bullet("Bearer authentication and HMAC signatures are optional tools; the receiving workflow must actually validate them.")
        self.add_bullet("Use the stable event identifier for idempotency where the receiver supports it.")
        self.add_bullet("Run Test webhook with a configuration-only synthetic object before enabling live delivery.")
        self.add_heading("Queue and failure behaviour", 2)
        self.add_bullet("Delivery runs away from the capture/decoder thread.")
        self.add_bullet("Transient failures retry with bounded exponential backoff; repeated failures move to DeadLetter for review.")
        self.add_bullet("Pending/dead-letter payloads may contain transformed decoded data and require the same protection as the destination.")
        self.add_callout("Third-party storage", "A webhook or hosting provider may retain requests, headers, files and event history. Do not send live decoded content until the service and workflow are approved for that data.", kind="warning")

    def chapter_data_outputs(self) -> None:
        self.add_chapter("12. Data outputs and delivery health", "data-outputs-and-delivery-health")
        self.add_figure(
            "07-data-outputs-v5.png",
            "Data outputs area in the v5 Settings family.",
            "Dark Settings screen with File transfer, Publish to web and Data outputs cards. No destination details or credentials are shown.",
            width=6.0,
        )
        self.add_heading("Privacy-aware optional adapters", 2)
        self.add_body("Optional adapters receive an immutable copy after PDW's established filtering and deduplication boundary. The output group and each adapter are disabled by default. Enabling the group requires an operator acknowledgement; filtered-only delivery, address masking and message omission are the privacy-safe defaults.")
        self.add_table(
            ["Adapter", "Use", "Important boundary"],
            [
                ("MQTT", "Publish JSON to a broker/topic.", "mqtts verifies certificate/hostname; plain mqtt needs explicit trusted-local opt-in; QoS 0 and no retain."),
                ("SQLite", "Write prepared, indexed local rows.", "Creates only when enabled/tested; event ID primary key makes replays idempotent."),
                ("MySQL / ODBC", "Write through a current Windows ODBC DSN.", "Driver, database and TLS policy are external; PDW uses prepared inserts."),
                ("Telnet JSON", "Read-only one-JSON-object-per-line stream.", "Loopback by default; non-loopback is explicit and unencrypted; no commands are accepted."),
                ("Windows", "Native per-user notifications.", "Message text separately opt-in because lock-screen display is possible."),
            ],
            [1800, 3000, 4560],
            font_size=7.9,
        )
        self.add_heading("Two-step routing gate", 2)
        self.add_body("A live message reaches an adapter only when the matching directory rule selects that destination and Settings has both the output group and adapter enabled/configured. Legacy non-directory-backed rules retain their documented compatibility behaviour. Configuration tests use synthetic events; SQLite/MySQL tests may deliberately write one identified test row.")
        self.add_heading("Queue and failure isolation", 2)
        self.add_body("Each adapter runs on a bounded worker queue of 500 events. An adapter error is recorded for health reporting and never blocks decoding or prevents another selected adapter from receiving its copy.")
        self.add_heading("Delivery Health", 2)
        self.add_body("Outputs > Delivery Health shows per-destination enabled state, successes, failures, drops, consecutive failures, last outcome and bounded recent operational history. Alerts may flash the taskbar or play a warning after a configured threshold without stealing focus.")
        self.add_callout("Content-free health", "Delivery Health stores destination names, timestamps, counters and sanitized summaries only. It never stores capcodes or decoded message text and cannot alter routing, retry, capture or decoding.", kind="success")

    def chapter_operations(self) -> None:
        self.add_chapter("13. Daily operation and troubleshooting", "daily-operation-and-troubleshooting")
        self.add_heading("Start-of-session checklist", 2)
        for item in (
            "Confirm the PDW version, architecture, instance/folder and selected profile.",
            "Confirm the intended source, decoder family and live meter without exposing private content.",
            "Check system date/time, disk space and message-history retention state.",
            "Confirm the correct Capcode Directory rules and whether Filter versus external routing is intended.",
            "Review Delivery Health and any publishing/transfer dead-letter items before enabling new delivery.",
            "Run a synthetic filtered and unfiltered test after a material configuration change.",
        ):
            self.add_bullet(item)
        self.add_heading("Encrypted configuration backup", 2)
        self.add_figure(
            "04-backup-restore-v5.png",
            "Password-encrypted Backup / Restore dialog (content-free native capture).",
            "Dark PDW Backup and Restore dialog explaining the complete configuration backup, two action buttons and a ready status. No path, password or operator data is shown.",
            width=5.4,
        )
        self.add_body("Settings > General > Backup / Restore creates one .pdwbackup file containing saved PDW.INI values, the complete Capcode Directory/filter rule set and supported Credential Manager records. AES-256-GCM protects the file with a user-chosen password derived using PBKDF2-HMAC-SHA256. PDW cannot recover a forgotten password.")
        self.add_bullet("Restore authenticates and validates the complete file before confirmation and transactional replacement.")
        self.add_bullet("A successful restore closes PDW so every restored setting loads together on restart.")
        self.add_bullet("The backup does not contain decoded messages, logs, recordings, output queues, website files, receiver DLLs or program binaries.")
        self.add_callout("Backup handling", "A .pdwbackup can contain usernames and secrets. Store the file and its password separately, restrict access, test restore in an isolated copy, and securely delete superseded backups under policy.", kind="warning")
        self.add_heading("Troubleshooting", 2)
        self.add_table(
            ["Symptom", "Checks"],
            [
                ("No live meter", "Confirm exact source, Windows permission, cable/device, saved endpoint identity, COM/receiver selection and source test."),
                ("Activity, no messages", "Confirm protocol family, frequency, PPM, rate, bandwidth, gain/level, antenna and clipping."),
                ("Named endpoint missing", "Reconnect/enable the exact VB-Cable endpoint or deliberately choose another input; PDW should not capture the default microphone."),
                ("RTL-SDR not listed", "Close other SDR software, confirm the exact USB interface/driver, and use a matching-bitness trusted receiver library."),
                ("Rule does not match", "Check protocol/capcode, exact versus + keywords, precedence, enabled state, Reject and Monitor Only."),
                ("No external delivery", "Confirm both the rule destination and Settings destination/group are enabled, configured and passing a synthetic test."),
                ("History CSV seems short", "Check Search, Protocol and Filtered controls; blank Message means text was not retained."),
                ("Queue grows", "Pause exposure if needed, correct the destination, review Delivery Health/dead letter, then retest with synthetic data."),
                ("Settings layout clips", "Record exact beta, architecture, Windows build, DPI, theme and content-free screenshot; do not attach operator configuration."),
            ],
            [2500, 6860],
            font_size=7.8,
        )
        self.add_heading("Safe acceptance", 2)
        self.add_steps((
            "Use an isolated training folder or test installation.",
            "Use synthetic, redacted or explicitly licensed recordings and directory rows.",
            "Exercise one known matching and one non-matching event through display, history and local outputs.",
            "Use disposable external test destinations and keep message text omitted.",
            "Confirm wrong TLS, authentication and SFTP host-key conditions fail closed.",
            "Record the exact beta, architecture and Windows environment without endpoint IDs, credentials or private traffic.",
            "Remove disposable credentials, exports and test rows after acceptance.",
        ))

    def chapter_reference(self) -> None:
        self.add_chapter("14. Reference", "reference")
        self.add_heading("Quick menu reference", 2)
        self.add_table(
            ["Menu", "Use"],
            [
                ("File", "Local logs/copies, configuration backup and exit."),
                ("Monitor", "Source, pause, mode, input/radio/replay, statistics and volume."),
                ("Filters", "Capcode Directory, rule behaviour, counters and migration compatibility."),
                ("Outputs", "Email, notification, transfer, publishing, adapters, health, dashboard and extra channels."),
                ("View", "Message history, columns, display and appearance."),
                ("Help", "Manual, diagnostics, About, licence and project resources."),
            ],
            [1800, 7560],
        )
        self.add_heading("Safe defaults", 2)
        self.add_table(
            ["Capability", "Default"],
            [
                ("Network and external outputs", "Disabled until deliberately configured and enabled."),
                ("Message history", "Disabled; storing message text separately disabled."),
                ("Local dashboard", "Disabled; loopback-only when enabled."),
                ("Guarded extra channels", "Disabled; each slot must be enabled and started."),
                ("FLEX fragment assembly", "Disabled in Public Beta 2."),
                ("Enhanced RTL IQ filtering", "Disabled; established RTL signal path retained."),
                ("Adelaide FLEX profile", "Standard settings remain the default; named profile is explicit and reversible."),
            ],
            [3200, 6160],
            font_size=8.3,
        )
        self.add_heading("Architecture and file reference", 2)
        self.add_table(
            ["Item", "Reference"],
            [
                ("Executable", "PDW v5.5.1 2026 Release.exe"),
                ("x64 role", "Current Windows and matching x64 libraries; RTL-TCP is architecture-neutral."),
                ("Win32 role", "Bundled x86 RTL-SDR and legacy 32-bit integrations."),
                ("Configuration", "PDW.INI beside the executable."),
                ("Directory/history", "pdw-history.sqlite3 by default."),
                ("Configuration export", "Password-encrypted .pdwbackup."),
                ("History export", "UTF-8 CSV outside archive retention."),
            ],
            [2700, 6660],
        )
        self.add_heading("Glossary", 2)
        self.add_table(
            ["Term", "Meaning"],
            [
                ("Capcode / Riccode", "Raw paging address used for matching; an operator alias never replaces it."),
                ("Filtered pane", "Lower pane selected by a rule's Filter action; independent from external delivery."),
                ("Monitor Only", "Keeps a match in the upper pane and does not make it a lower-pane filtered result."),
                ("WASAPI", "Modern Windows audio capture API used for exact saved endpoints and as a controlled fallback."),
                ("WinMM", "Established Windows multimedia audio path retained for compatibility."),
                ("RTL-TCP", "TCP service supplying tuner IQ samples; no decoded message text is sent to it."),
                ("SigMF", "Signal Metadata Format used here for real float32 diagnostic samples plus metadata."),
                ("FTPS / SFTP", "TLS-protected FTP / SSH File Transfer Protocol; they use different verification models."),
                ("Dead letter", "Repeatedly failed queued delivery retained for authorised operator review."),
                ("Development preview", "Current source work that is not yet part of the published beta contract."),
            ],
            [2400, 6960],
            font_size=8.0,
        )
        self.add_heading("Credits and licence", 2)
        self.add_body("PDW retains credits to Jason Petty, Peter Hunt, the 2013 open-source release and later contributors, including the current modernization work credited by the application. The application and source identify GNU General Public License version 3 or later. Review the supplied License and THIRD_PARTY_NOTICES.md for complete terms and component notices.")
        self.add_callout("Manual maintenance", "This edition targets the immutable v5.5.1 Public Beta 2 release and separately identifies 11 August 2026 development previews. Re-check version, menus, security guidance, third-party service requirements and hardware acceptance before packaging it with a later beta or stable release.", kind="note")

    def set_core_properties(self) -> None:
        properties = self.document.core_properties
        properties.title = "PDW v5.5.1 2026 Release User Manual"
        properties.subject = "Operator guide for PDW Paging Decoder for Windows - Public Beta 2"
        properties.author = "PDW Community Project"
        properties.last_modified_by = "PDW Community Project"
        edition_timestamp = datetime(2026, 8, 11, 13, 0, 0, tzinfo=timezone.utc)
        properties.created = edition_timestamp
        properties.modified = edition_timestamp
        properties.keywords = "PDW, POCSAG, FLEX, ACARS, MOBITEX, ERMES, RTL-SDR, Apprise, Capcode Directory"
        properties.comments = "All examples are synthetic or masked training data."


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the PDW v5.5.1 illustrated user manual.")
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    builder = V551ManualBuilder(args.output.resolve())
    builder.build()
    builder.save()
    print(f"Wrote {builder.output}")


if __name__ == "__main__":
    main()
